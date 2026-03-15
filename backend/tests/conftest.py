from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from app.core.config import get_settings
from app.services.scenario_service import scenario_service


@pytest.fixture
def fixtures_dir() -> Path:
    return Path(__file__).resolve().parent / "fixtures"


@pytest.fixture
def sample_prompt(fixtures_dir: Path) -> str:
    return (fixtures_dir / "chat_prompt.txt").read_text(encoding="utf-8").strip()


@pytest.fixture
def docx_prompt(fixtures_dir: Path) -> str:
    return (fixtures_dir / "docx_prompt_example.txt").read_text(encoding="utf-8").strip()


@pytest.fixture
def docx_fixtures_dir(fixtures_dir: Path) -> Path:
    return fixtures_dir / "docx"


@pytest.fixture
def sample_docx_template_path(docx_fixtures_dir: Path) -> Path:
    return docx_fixtures_dir / "sample_template.docx"


@pytest.fixture
def sample_docx_good_example_path(docx_fixtures_dir: Path) -> Path:
    return docx_fixtures_dir / "sample_good_example.docx"


@pytest.fixture
def sample_docx_bad_example_path(docx_fixtures_dir: Path) -> Path:
    return docx_fixtures_dir / "sample_bad_example.docx"


@pytest.fixture
def sample_docx_enhancement_with_image_path(docx_fixtures_dir: Path) -> Path:
    return docx_fixtures_dir / "sample_enhancement_with_image.docx"


@pytest.fixture
def expected_docx_output_path(docx_fixtures_dir: Path) -> Path:
    return docx_fixtures_dir / "expected_final_output.docx"


@pytest.fixture
def sample_template_path(tmp_path: Path) -> Path:
    path = tmp_path / "template.docx"
    document = Document()
    document.add_heading("Project Overview", level=1)
    document.add_paragraph("Describe the high-level objective.")
    document.add_heading("Functional Requirements", level=1)
    document.add_paragraph("List required business capabilities.")
    document.save(str(path))
    return path


@pytest.fixture
def isolated_storage(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> Path:
    settings = get_settings()
    storage_root = tmp_path / "storage"
    scenarios_root = tmp_path / "scenarios"
    monkeypatch.setattr(settings, "storage_root", storage_root)
    monkeypatch.setattr(settings, "vector_store_root", storage_root / "vectorstores")
    monkeypatch.setattr(settings, "upload_root", storage_root / "uploads")
    monkeypatch.setattr(settings, "generated_root", storage_root / "generated")
    monkeypatch.setattr(settings, "scenarios_root", scenarios_root)
    monkeypatch.setattr(settings, "scenarios_db_path", scenarios_root / "scenarios.db")
    settings.storage_root.mkdir(parents=True, exist_ok=True)
    settings.vector_store_root.mkdir(parents=True, exist_ok=True)
    settings.upload_root.mkdir(parents=True, exist_ok=True)
    settings.generated_root.mkdir(parents=True, exist_ok=True)
    settings.scenarios_root.mkdir(parents=True, exist_ok=True)
    scenario_service._ensure_database()
    return tmp_path