from __future__ import annotations

import asyncio
from contextlib import asynccontextmanager
import sys
import zipfile
from pathlib import Path
import json

from docx import Document
from fastapi.testclient import TestClient
import pytest

ROOT_DIR = Path(__file__).resolve().parents[1]
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

from app.main import app
from app.models.document_state import McpServerCatalogResponse, McpServerSummary
from app.services.docker_mcp_service import DockerMcpToolCallResult, DockerMcpToolSpec, docker_mcp_service
import app.services.llm_provider as llm_provider_module
from app.services.llm_provider import llm_provider
from app.services.rag_service import rag_service
from app.services.session_store import session_store
from app.services.translation_service import translation_service
from scripts.run_batch_workflow import run_batch_workflow


def test_llm_provider_includes_selected_mcp_servers_in_system_prompt(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(
        docker_mcp_service,
        "list_servers",
        lambda: McpServerCatalogResponse(
            available=True,
            servers=[McpServerSummary(name="fetch", description="Fetches a URL from the internet.")],
        ),
    )

    prompt = llm_provider.with_scenario_mcp_context("Base prompt.", ["fetch"])

    assert "Base prompt." in prompt
    assert "Scenario MCP server context" in prompt
    assert "fetch" in prompt
    assert "Fetches a URL from the internet." in prompt


def test_docker_mcp_service_uses_dedicated_sse_read_timeout(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(
        docker_mcp_service,
        "settings",
        type(
            "SettingsStub",
            (),
            {"request_timeout_seconds": 45.0, "docker_mcp_sse_read_timeout_seconds": 3600.0},
        )(),
    )

    assert docker_mcp_service._get_sse_read_timeout_seconds() == 3600.0

    monkeypatch.setattr(
        docker_mcp_service,
        "settings",
        type(
            "SettingsStub",
            (),
            {"request_timeout_seconds": 7200.0, "docker_mcp_sse_read_timeout_seconds": 3600.0},
        )(),
    )

    assert docker_mcp_service._get_sse_read_timeout_seconds() == 7200.0


def test_llm_provider_executes_selected_mcp_tools(monkeypatch: pytest.MonkeyPatch) -> None:
    payloads: list[dict[str, object]] = []
    tool_calls: list[tuple[str, dict[str, object]]] = []

    class FakeResponse:
        def __init__(self, payload: dict[str, object]) -> None:
            self._payload = payload
            self.text = json.dumps(payload)

        def raise_for_status(self) -> None:
            return None

        def json(self) -> dict[str, object]:
            return self._payload

    responses = [
        FakeResponse(
            {
                "choices": [
                    {
                        "message": {
                            "content": "",
                            "tool_calls": [
                                {
                                    "id": "call-fetch-1",
                                    "type": "function",
                                    "function": {
                                        "name": "fetch",
                                        "arguments": '{"url": "https://www.jn.pt"}',
                                    },
                                }
                            ],
                        }
                    }
                ]
            }
        ),
        FakeResponse({"choices": [{"message": {"content": "Fetched content summary."}}]}),
    ]

    class FakeAsyncClient:
        def __init__(self, *args, **kwargs) -> None:
            pass

        async def __aenter__(self) -> "FakeAsyncClient":
            return self

        async def __aexit__(self, exc_type, exc, tb) -> None:
            return None

        async def post(self, url: str, json: dict[str, object], headers: dict[str, str]) -> FakeResponse:
            payloads.append(json)
            return responses.pop(0)

    class FakeGatewayClient:
        async def list_tools(self) -> list[DockerMcpToolSpec]:
            return [
                DockerMcpToolSpec(
                    name="fetch",
                    description="Fetch a URL from the internet.",
                    input_schema={
                        "type": "object",
                        "properties": {"url": {"type": "string"}},
                        "required": ["url"],
                    },
                )
            ]

        async def call_tool(self, tool_name: str, arguments: dict[str, object]) -> DockerMcpToolCallResult:
            tool_calls.append((tool_name, arguments))
            return DockerMcpToolCallResult(
                text="{\"url\": \"https://www.jn.pt\", \"content\": \"Headline from JN\"}",
                structured_content={"url": "https://www.jn.pt", "content": "Headline from JN"},
            )

    @asynccontextmanager
    async def fake_tool_client(server_names: list[str]):
        assert server_names == ["fetch"]
        yield FakeGatewayClient()

    monkeypatch.setattr(docker_mcp_service, "tool_client", fake_tool_client)
    monkeypatch.setattr(llm_provider_module.httpx, "AsyncClient", FakeAsyncClient)

    result = asyncio.run(
        llm_provider.chat_completion(
            system_prompt="Use tools when needed.",
            user_prompt="Load data from jn.pt.",
            mcp_servers=["fetch"],
        )
    )

    assert result == "Fetched content summary."
    assert tool_calls == [("fetch", {"url": "https://www.jn.pt"})]
    assert payloads[0]["tool_choice"] == "auto"
    assert payloads[0]["tools"][0]["function"]["name"] == "fetch"
    second_messages = payloads[1]["messages"]
    assert any(message.get("role") == "tool" for message in second_messages)


def test_llm_provider_generate_json_extracts_embedded_json_prefix(monkeypatch: pytest.MonkeyPatch) -> None:
    async def fake_chat_completion(**kwargs) -> str:
        return (
            "I've fetched the content from www.sic.pt and will update the draft now: "
            '{"assistant_message":"Fetched sic.pt.","summary":"sic.pt summary.","section_updates":[]}'
        )

    monkeypatch.setattr(llm_provider, "chat_completion", fake_chat_completion)

    payload = asyncio.run(
        llm_provider.generate_json(
            system_prompt="Return JSON.",
            user_prompt="Fetch sic.pt.",
        )
    )

    assert payload["assistant_message"] == "Fetched sic.pt."
    assert payload["summary"] == "sic.pt summary."
    assert payload["section_updates"] == []


def test_llm_provider_forces_finalization_after_repeated_tool_call(monkeypatch: pytest.MonkeyPatch) -> None:
    payloads: list[dict[str, object]] = []

    class FakeResponse:
        def __init__(self, payload: dict[str, object]) -> None:
            self._payload = payload
            self.text = json.dumps(payload)

        def raise_for_status(self) -> None:
            return None

        def json(self) -> dict[str, object]:
            return self._payload

    responses = [
        FakeResponse(
            {
                "choices": [
                    {
                        "message": {
                            "content": "",
                            "tool_calls": [
                                {
                                    "id": "call-fetch-1",
                                    "type": "function",
                                    "function": {"name": "fetch", "arguments": '{"url": "https://www.sic.pt"}'},
                                }
                            ],
                        }
                    }
                ]
            }
        ),
        FakeResponse(
            {
                "choices": [
                    {
                        "message": {
                            "content": "",
                            "tool_calls": [
                                {
                                    "id": "call-fetch-2",
                                    "type": "function",
                                    "function": {"name": "fetch", "arguments": '{"url": "https://www.sic.pt"}'},
                                }
                            ],
                        }
                    }
                ]
            }
        ),
        FakeResponse({"choices": [{"message": {"content": "Finalized after tool reuse."}}]}),
    ]

    class FakeAsyncClient:
        def __init__(self, *args, **kwargs) -> None:
            pass

        async def __aenter__(self) -> "FakeAsyncClient":
            return self

        async def __aexit__(self, exc_type, exc, tb) -> None:
            return None

        async def post(self, url: str, json: dict[str, object], headers: dict[str, str]) -> FakeResponse:
            payloads.append(json)
            return responses.pop(0)

    class FakeGatewayClient:
        async def list_tools(self) -> list[DockerMcpToolSpec]:
            return [
                DockerMcpToolSpec(
                    name="fetch",
                    description="Fetch a URL from the internet.",
                    input_schema={"type": "object", "properties": {"url": {"type": "string"}}, "required": ["url"]},
                )
            ]

        async def call_tool(self, tool_name: str, arguments: dict[str, object]) -> DockerMcpToolCallResult:
            return DockerMcpToolCallResult(text="Fetched sic.pt content")

    @asynccontextmanager
    async def fake_tool_client(server_names: list[str]):
        yield FakeGatewayClient()

    monkeypatch.setattr(docker_mcp_service, "tool_client", fake_tool_client)
    monkeypatch.setattr(llm_provider_module.httpx, "AsyncClient", FakeAsyncClient)

    result = asyncio.run(
        llm_provider.chat_completion(
            system_prompt="Use tools when needed.",
            user_prompt="Load data from sic.pt.",
            mcp_servers=["fetch"],
        )
    )

    assert result == "Finalized after tool reuse."
    assert payloads[0]["tool_choice"] == "auto"
    assert payloads[1]["tool_choice"] == "auto"
    assert payloads[2]["tool_choice"] == "none"


def test_llm_provider_executes_inline_tool_markup(monkeypatch: pytest.MonkeyPatch) -> None:
    payloads: list[dict[str, object]] = []
    tool_calls: list[tuple[str, dict[str, object]]] = []

    class FakeResponse:
        def __init__(self, payload: dict[str, object]) -> None:
            self._payload = payload
            self.text = json.dumps(payload)

        def raise_for_status(self) -> None:
            return None

        def json(self) -> dict[str, object]:
            return self._payload

    responses = [
        FakeResponse(
            {
                "choices": [
                    {
                        "message": {
                            "content": (
                                "I need to update the data retention content to include information from www.rtp.pt. "
                                "Let me fetch that content first to understand what specific information should be included. "
                                "<function=fetch> <parameter=url> http://www.rtp.pt </parameter> </function>"
                            )
                        }
                    }
                ]
            }
        ),
        FakeResponse(
            {
                "choices": [
                    {
                        "message": {
                            "content": '{"assistant_message":"Fetched RTP.","summary":"rtp.pt summary.","section_updates":[]}'
                        }
                    }
                ]
            }
        ),
    ]

    class FakeAsyncClient:
        def __init__(self, *args, **kwargs) -> None:
            pass

        async def __aenter__(self) -> "FakeAsyncClient":
            return self

        async def __aexit__(self, exc_type, exc, tb) -> None:
            return None

        async def post(self, url: str, json: dict[str, object], headers: dict[str, str]) -> FakeResponse:
            payloads.append(json)
            return responses.pop(0)

    class FakeGatewayClient:
        async def list_tools(self) -> list[DockerMcpToolSpec]:
            return [
                DockerMcpToolSpec(
                    name="fetch",
                    description="Fetch a URL from the internet.",
                    input_schema={"type": "object", "properties": {"url": {"type": "string"}}, "required": ["url"]},
                )
            ]

        async def call_tool(self, tool_name: str, arguments: dict[str, object]) -> DockerMcpToolCallResult:
            tool_calls.append((tool_name, arguments))
            return DockerMcpToolCallResult(text="Fetched RTP content")

    @asynccontextmanager
    async def fake_tool_client(server_names: list[str]):
        yield FakeGatewayClient()

    monkeypatch.setattr(docker_mcp_service, "tool_client", fake_tool_client)
    monkeypatch.setattr(llm_provider_module.httpx, "AsyncClient", FakeAsyncClient)

    payload = asyncio.run(
        llm_provider.generate_json(
            system_prompt="Return JSON.",
            user_prompt="Fetch rtp.pt.",
            mcp_servers=["fetch"],
        )
    )

    assert payload["assistant_message"] == "Fetched RTP."
    assert tool_calls == [("fetch", {"url": "http://www.rtp.pt"})]
    assert any(message.get("role") == "tool" for message in payloads[1]["messages"])


def test_llm_provider_ignores_tool_calls_after_tool_usage_disabled(monkeypatch: pytest.MonkeyPatch) -> None:
    payloads: list[dict[str, object]] = []
    executed_calls: list[tuple[str, dict[str, object]]] = []

    class FakeResponse:
        def __init__(self, payload: dict[str, object]) -> None:
            self._payload = payload
            self.text = json.dumps(payload)

        def raise_for_status(self) -> None:
            return None

        def json(self) -> dict[str, object]:
            return self._payload

    responses = [
        FakeResponse(
            {
                "choices": [
                    {
                        "message": {
                            "content": "",
                            "tool_calls": [
                                {
                                    "id": "call-1",
                                    "type": "function",
                                    "function": {"name": "fetch", "arguments": '{"url": "https://a.example"}'},
                                }
                            ],
                        }
                    }
                ]
            }
        ),
        FakeResponse(
            {
                "choices": [
                    {
                        "message": {
                            "content": "",
                            "tool_calls": [
                                {
                                    "id": "call-2",
                                    "type": "function",
                                    "function": {"name": "fetch", "arguments": '{"url": "https://b.example"}'},
                                }
                            ],
                        }
                    }
                ]
            }
        ),
        FakeResponse(
            {
                "choices": [
                    {
                        "message": {
                            "content": "",
                            "tool_calls": [
                                {
                                    "id": "call-3",
                                    "type": "function",
                                    "function": {"name": "fetch", "arguments": '{"url": "https://c.example"}'},
                                }
                            ],
                        }
                    }
                ]
            }
        ),
        FakeResponse(
            {
                "choices": [
                    {
                        "message": {
                            "content": "Still trying tools",
                            "tool_calls": [
                                {
                                    "id": "call-4",
                                    "type": "function",
                                    "function": {"name": "fetch", "arguments": '{"url": "https://d.example"}'},
                                }
                            ],
                        }
                    }
                ]
            }
        ),
        FakeResponse({"choices": [{"message": {"content": "Final answer without more tool calls."}}]}),
    ]

    class FakeAsyncClient:
        def __init__(self, *args, **kwargs) -> None:
            pass

        async def __aenter__(self) -> "FakeAsyncClient":
            return self

        async def __aexit__(self, exc_type, exc, tb) -> None:
            return None

        async def post(self, url: str, json: dict[str, object], headers: dict[str, str]) -> FakeResponse:
            payloads.append(json)
            return responses.pop(0)

    class FakeGatewayClient:
        async def list_tools(self) -> list[DockerMcpToolSpec]:
            return [
                DockerMcpToolSpec(
                    name="fetch",
                    description="Fetch a URL from the internet.",
                    input_schema={"type": "object", "properties": {"url": {"type": "string"}}, "required": ["url"]},
                )
            ]

        async def call_tool(self, tool_name: str, arguments: dict[str, object]) -> DockerMcpToolCallResult:
            executed_calls.append((tool_name, arguments))
            return DockerMcpToolCallResult(text=f"Fetched {arguments['url']}")

    @asynccontextmanager
    async def fake_tool_client(server_names: list[str]):
        yield FakeGatewayClient()

    monkeypatch.setattr(docker_mcp_service, "tool_client", fake_tool_client)
    monkeypatch.setattr(llm_provider_module.httpx, "AsyncClient", FakeAsyncClient)

    result = asyncio.run(
        llm_provider.chat_completion(
            system_prompt="Use tools when needed.",
            user_prompt="Gather data.",
            mcp_servers=["fetch"],
        )
    )

    assert result == "Final answer without more tool calls."
    assert executed_calls == [
        ("fetch", {"url": "https://a.example"}),
        ("fetch", {"url": "https://b.example"}),
        ("fetch", {"url": "https://c.example"}),
    ]
    assert payloads[3]["tool_choice"] == "none"


def test_batch_workflow_generates_localized_archive(
    tmp_path,
    monkeypatch: pytest.MonkeyPatch,
    sample_template_path: Path,
    sample_prompt: str,
    isolated_storage: Path,
):
    template_path = sample_template_path
    good_example_path = ROOT_DIR / "tests" / "fixtures" / "good_example.txt"
    bad_example_path = ROOT_DIR / "tests" / "fixtures" / "bad_example.txt"

    monkeypatch.setattr(rag_service, "index_examples", lambda *args, **kwargs: [])

    async def fake_generate_json(*, system_prompt: str, user_prompt: str, temperature: float = 0.2, mcp_servers=None):
        if "assistant_message" in user_prompt:
            return {
                "assistant_message": "Captured the core scope. Please confirm any audit retention rules.",
                "summary": "Compliance onboarding workflow for enterprise customers.",
                "section_updates": [
                    {
                        "section_id": "section-1",
                        "title": "Project Overview",
                        "content": "Build a compliance onboarding workflow for enterprise customers.",
                        "status": "complete",
                    },
                    {
                        "section_id": "section-2",
                        "title": "Functional Requirements",
                        "content": "Include identity verification, document review, approval audit trail, and reviewer notifications.",
                        "status": "complete",
                    },
                ],
            }

        if "Target language: Spanish." in system_prompt:
            return {
                "sections": [
                    {
                        "section_id": "section-1",
                        "content": "Crear un flujo de incorporacion de cumplimiento para clientes empresariales.",
                    },
                    {
                        "section_id": "section-2",
                        "content": "Incluir verificacion de identidad, revision documental, trazabilidad de aprobacion y notificaciones.",
                    },
                ]
            }

        if "Target language: French." in system_prompt:
            return {
                "sections": [
                    {
                        "section_id": "section-1",
                        "content": "Creer un flux d'integration de conformite pour les clients entreprise.",
                    },
                    {
                        "section_id": "section-2",
                        "content": "Inclure la verification d'identite, la revue documentaire, la piste d'audit et les notifications.",
                    },
                ]
            }

        raise AssertionError(f"Unexpected LLM prompt: {system_prompt}")

    monkeypatch.setattr(llm_provider, "generate_json", fake_generate_json)

    with TestClient(app) as client:
        result = run_batch_workflow(
            base_url="http://testserver",
            template_path=template_path,
            good_example_paths=[good_example_path],
            bad_example_paths=[bad_example_path],
            message=sample_prompt,
            languages=["Spanish", "French"],
            client=client,
        )

    ingest_payload = result["ingest"]
    chat_payload = result["chat"]
    export_payload = result["export"]

    assert ingest_payload["template"]["file_type"] == "docx"
    assert len(ingest_payload["template"]["sections"]) == 2
    assert chat_payload["llm_available"] is True
    assert "Compliance onboarding workflow" in chat_payload["preview_markdown"]

    archive_path = Path(export_payload["archive_path"])
    assert archive_path.exists()
    assert len(export_payload["generated_files"]) == 2

    with zipfile.ZipFile(archive_path) as archive:
        archive_names = set(archive.namelist())
        assert "template.spanish.docx" in archive_names
        assert "template.french.docx" in archive_names

    spanish_document = Document(str(archive_path.parent / "template.spanish.docx"))
    spanish_text = "\n".join(paragraph.text for paragraph in spanish_document.paragraphs)
    assert "Crear un flujo de incorporacion de cumplimiento" in spanish_text

    session_id = ingest_payload["session_id"]
    assert session_store.get(session_id).draft_state.sections[0].status == "complete"


def test_chat_updates_section_when_llm_returns_title_instead_of_section_id(
    monkeypatch: pytest.MonkeyPatch,
    sample_template_path: Path,
    isolated_storage: Path,
) -> None:
    monkeypatch.setattr(rag_service, "index_examples", lambda *args, **kwargs: [])

    async def fake_generate_json(*, system_prompt: str, user_prompt: str, temperature: float = 0.2, mcp_servers=None):
        if '"assistant_message"' in user_prompt:
            return {
                "assistant_message": "Updated the requested field.",
                "summary": "One field updated.",
                "section_updates": [
                    {
                        "section_id": "project overview",
                        "title": "Project Overview",
                        "content": "Nuno Pereira",
                        "status": "complete",
                    }
                ],
            }
        raise AssertionError(f"Unexpected prompt: {system_prompt}")

    monkeypatch.setattr(llm_provider, "generate_json", fake_generate_json)

    with TestClient(app) as client:
        with sample_template_path.open("rb") as template_handle:
            ingest_response = client.post(
                "/ingest",
                files={
                    "template": (
                        sample_template_path.name,
                        template_handle.read(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )
        ingest_response.raise_for_status()
        session_id = ingest_response.json()["session_id"]

        chat_response = client.post(
            "/chat",
            json={"session_id": session_id, "message": "Set Project Overview to Nuno Pereira."},
        )
        chat_response.raise_for_status()

    payload = chat_response.json()
    assert "Nuno Pereira" in payload["preview_markdown"]
    assert payload["draft_state"]["sections"][0]["content"] == "Nuno Pereira"


def test_chat_falls_back_to_user_assignment_when_llm_returns_no_section_updates(
    monkeypatch: pytest.MonkeyPatch,
    sample_docx_template_path: Path,
    sample_docx_good_example_path: Path,
    sample_docx_bad_example_path: Path,
    isolated_storage: Path,
) -> None:
    monkeypatch.setattr(rag_service, "index_examples", lambda *args, **kwargs: [])

    async def fake_generate_json(*, system_prompt: str, user_prompt: str, temperature: float = 0.2, mcp_servers=None):
        if '"assistant_message"' in user_prompt:
            return {
                "assistant_message": "Updating the Project Overview field with the provided name.",
                "summary": "Project Overview section has been updated with 'Nuno Pereira'.",
                "section_updates": [],
            }
        raise AssertionError(f"Unexpected prompt: {system_prompt}")

    monkeypatch.setattr(llm_provider, "generate_json", fake_generate_json)

    with TestClient(app) as client:
        with sample_docx_template_path.open("rb") as template_handle, sample_docx_good_example_path.open("rb") as good_handle, sample_docx_bad_example_path.open("rb") as bad_handle:
            ingest_response = client.post(
                "/ingest",
                files=[
                    (
                        "template",
                        (
                            sample_docx_template_path.name,
                            template_handle.read(),
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        ),
                    ),
                    (
                        "good_examples",
                        (
                            sample_docx_good_example_path.name,
                            good_handle.read(),
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        ),
                    ),
                    (
                        "bad_examples",
                        (
                            sample_docx_bad_example_path.name,
                            bad_handle.read(),
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        ),
                    ),
                ],
            )
        ingest_response.raise_for_status()
        session_id = ingest_response.json()["session_id"]

        chat_response = client.post(
            "/chat",
            json={"session_id": session_id, "message": 'fill "Nuno Pereira" on "Project Overview" field'},
        )
        chat_response.raise_for_status()

    payload = chat_response.json()
    assert payload["draft_state"]["sections"][0]["content"] == "Nuno Pereira"
    assert "Nuno Pereira" in payload["preview_markdown"]


def test_chat_can_append_new_section_when_llm_returns_a_new_title(
    monkeypatch: pytest.MonkeyPatch,
    sample_template_path: Path,
    isolated_storage: Path,
) -> None:
    monkeypatch.setattr(rag_service, "index_examples", lambda *args, **kwargs: [])

    async def fake_generate_json(*, system_prompt: str, user_prompt: str, temperature: float = 0.2, mcp_servers=None):
        if '"assistant_message"' in user_prompt:
            return {
                "assistant_message": "Added the Conclusion section.",
                "summary": "A new conclusion section was added.",
                "section_updates": [
                    {
                        "title": "Conclusion",
                        "content": "2026-03-15 18:12:00 WET",
                        "status": "complete",
                    }
                ],
            }
        raise AssertionError(f"Unexpected prompt: {system_prompt}")

    monkeypatch.setattr(llm_provider, "generate_json", fake_generate_json)

    with TestClient(app) as client:
        with sample_template_path.open("rb") as template_handle:
            ingest_response = client.post(
                "/ingest",
                files={
                    "template": (
                        sample_template_path.name,
                        template_handle.read(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )
        ingest_response.raise_for_status()
        session_id = ingest_response.json()["session_id"]

        chat_response = client.post(
            "/chat",
            json={"session_id": session_id, "message": "Add a new section named Conclusion with the date and time of today."},
        )
        chat_response.raise_for_status()

    payload = chat_response.json()
    dynamic_section = next(section for section in payload["draft_state"]["sections"] if section["title"] == "Conclusion")
    assert dynamic_section["content"] == "2026-03-15 18:12:00 WET"
    assert "## Conclusion" in payload["preview_markdown"]


def test_chat_can_append_new_section_from_user_request_when_llm_returns_no_updates(
    monkeypatch: pytest.MonkeyPatch,
    sample_template_path: Path,
    isolated_storage: Path,
) -> None:
    monkeypatch.setattr(rag_service, "index_examples", lambda *args, **kwargs: [])

    async def fake_generate_json(*, system_prompt: str, user_prompt: str, temperature: float = 0.2, mcp_servers=None):
        if '"assistant_message"' in user_prompt:
            return {
                "assistant_message": "Adding a new section named Conclusion with today's date and time.",
                "summary": "Conclusion section added.",
                "section_updates": [],
            }
        raise AssertionError(f"Unexpected prompt: {system_prompt}")

    monkeypatch.setattr(llm_provider, "generate_json", fake_generate_json)

    with TestClient(app) as client:
        with sample_template_path.open("rb") as template_handle:
            ingest_response = client.post(
                "/ingest",
                files={
                    "template": (
                        sample_template_path.name,
                        template_handle.read(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )
        ingest_response.raise_for_status()
        session_id = ingest_response.json()["session_id"]

        chat_response = client.post(
            "/chat",
            json={"session_id": session_id, "message": "Add a new section named Conclusion with the date and time of today."},
        )
        chat_response.raise_for_status()

    payload = chat_response.json()
    dynamic_section = next(section for section in payload["draft_state"]["sections"] if section["title"] == "Conclusion")
    assert dynamic_section["content"]
    assert "## Conclusion" in payload["preview_markdown"]


def test_chat_can_place_move_and_delete_sections_from_user_requests(
    monkeypatch: pytest.MonkeyPatch,
    sample_template_path: Path,
    isolated_storage: Path,
) -> None:
    monkeypatch.setattr(rag_service, "index_examples", lambda *args, **kwargs: [])

    async def fake_generate_json(*, system_prompt: str, user_prompt: str, temperature: float = 0.2, mcp_servers=None):
        if '"assistant_message"' in user_prompt:
            return {
                "assistant_message": "Applied the requested structural edit.",
                "summary": "Draft structure updated.",
                "section_updates": [],
                "section_operations": [],
            }
        raise AssertionError(f"Unexpected prompt: {system_prompt}")

    monkeypatch.setattr(llm_provider, "generate_json", fake_generate_json)

    with TestClient(app) as client:
        with sample_template_path.open("rb") as template_handle:
            ingest_response = client.post(
                "/ingest",
                files={
                    "template": (
                        sample_template_path.name,
                        template_handle.read(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )
        ingest_response.raise_for_status()
        session_id = ingest_response.json()["session_id"]

        add_response = client.post(
            "/chat",
            json={"session_id": session_id, "message": "Add a new section named Introduction with Saturday at the top."},
        )
        add_response.raise_for_status()

        move_response = client.post(
            "/chat",
            json={"session_id": session_id, "message": "Move the Functional Requirements section after the Introduction section."},
        )
        move_response.raise_for_status()

        delete_response = client.post(
            "/chat",
            json={"session_id": session_id, "message": "Delete the Project Overview section."},
        )
        delete_response.raise_for_status()

    payload = delete_response.json()
    section_titles = [section["title"] for section in payload["draft_state"]["sections"]]
    assert section_titles == ["Introduction", "Functional Requirements"]
    assert "## Introduction" in payload["preview_markdown"]
    assert "## Project Overview" not in payload["preview_markdown"]


def test_export_respects_structural_section_operations(
    monkeypatch: pytest.MonkeyPatch,
    sample_template_path: Path,
    isolated_storage: Path,
) -> None:
    monkeypatch.setattr(rag_service, "index_examples", lambda *args, **kwargs: [])

    async def fake_generate_json(*, system_prompt: str, user_prompt: str, temperature: float = 0.2, mcp_servers=None):
        if '"assistant_message"' in user_prompt:
            return {
                "assistant_message": "Applied the requested structural edit.",
                "summary": "Draft structure updated.",
                "section_updates": [],
                "section_operations": [
                    {
                        "action": "add",
                        "title": "Introduction",
                        "content": "Opening notes.",
                        "position": "top",
                        "status": "complete",
                    },
                    {
                        "action": "move",
                        "title": "Functional Requirements",
                        "position": "after",
                        "relative_title": "Introduction",
                    },
                    {
                        "action": "delete",
                        "title": "Project Overview",
                    },
                ],
            }
        raise AssertionError(f"Unexpected prompt: {system_prompt}")

    async def fake_translate_sections(*, session, language: str, good_examples):
        assert language == "Spanish"
        translated: dict[str, str] = {}
        for section in session.draft_state.sections:
            if section.title == "Introduction":
                translated[f"{section.section_id}::title"] = "Introduccion"
                translated[f"{section.section_id}::content"] = "Notas de apertura."
                translated[section.section_id] = "Notas de apertura."
        return translated

    monkeypatch.setattr(llm_provider, "generate_json", fake_generate_json)
    monkeypatch.setattr(translation_service, "translate_sections", fake_translate_sections)

    with TestClient(app) as client:
        with sample_template_path.open("rb") as template_handle:
            ingest_response = client.post(
                "/ingest",
                files={
                    "template": (
                        sample_template_path.name,
                        template_handle.read(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )
        ingest_response.raise_for_status()
        session_id = ingest_response.json()["session_id"]

        chat_response = client.post(
            "/chat",
            json={"session_id": session_id, "message": "Add Introduction to the top, move Functional Requirements after it, and delete Project Overview."},
        )
        chat_response.raise_for_status()

        export_response = client.post(
            "/export",
            json={"session_id": session_id, "target_languages": ["Spanish"], "output_file_name": "structured-spec"},
        )
        export_response.raise_for_status()

    exported_path = Path(export_response.json()["generated_files"][0])
    document = Document(str(exported_path))
    text_lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    assert text_lines[0] == "Introduccion"
    assert text_lines[1] == "Notas de apertura."
    assert "Project Overview" not in text_lines
    assert "Functional Requirements" in text_lines


def test_export_can_recover_session_after_in_memory_store_is_cleared(
    monkeypatch: pytest.MonkeyPatch,
    sample_template_path: Path,
    isolated_storage: Path,
) -> None:
    monkeypatch.setattr(rag_service, "index_examples", lambda *args, **kwargs: [])

    async def fake_generate_json(*, system_prompt: str, user_prompt: str, temperature: float = 0.2, mcp_servers=None):
        if '"assistant_message"' in user_prompt:
            return {
                "assistant_message": "Captured the scope.",
                "summary": "Draft ready.",
                "section_updates": [
                    {
                        "section_id": "section-1",
                        "title": "Project Overview",
                        "content": "Persistent export scope.",
                        "status": "complete",
                    }
                ],
            }
        raise AssertionError(f"Unexpected prompt: {system_prompt}")

    async def fake_translate_sections(*, session, language: str, good_examples):
        return {
            "section-1::title": "Resumen del proyecto",
            "section-1::content": "Alcance persistente de exportacion.",
            "section-1": "Alcance persistente de exportacion.",
        }

    monkeypatch.setattr(llm_provider, "generate_json", fake_generate_json)
    monkeypatch.setattr(translation_service, "translate_sections", fake_translate_sections)

    with TestClient(app) as client:
        with sample_template_path.open("rb") as template_handle:
            ingest_response = client.post(
                "/ingest",
                files={
                    "template": (
                        sample_template_path.name,
                        template_handle.read(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )
        ingest_response.raise_for_status()
        session_id = ingest_response.json()["session_id"]

        chat_response = client.post(
            "/chat",
            json={"session_id": session_id, "message": "Set Project Overview to Persistent export scope."},
        )
        chat_response.raise_for_status()

        session_store._sessions.clear()

        export_response = client.post(
            "/export",
            json={"session_id": session_id, "target_languages": ["Spanish"], "output_file_name": "persistent-spec"},
        )
        export_response.raise_for_status()

    exported_path = Path(export_response.json()["generated_files"][0])
    assert exported_path.exists()


def test_translation_configuration_endpoint_reports_active_provider(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(translation_service.settings, "translation_provider", "google")
    monkeypatch.setattr(translation_service.settings, "google_translate_api_key", "fake-key")
    monkeypatch.setattr(translation_service.settings, "azure_translator_key", None)
    monkeypatch.setattr(translation_service.settings, "azure_translator_region", None)

    with TestClient(app) as client:
        response = client.get("/config/translation")

    response.raise_for_status()
    payload = response.json()
    assert payload["active_provider"] == "google"
    options = {option["id"]: option for option in payload["options"]}
    assert options["google"]["configured"] is True
    assert options["azure"]["configured"] is False


def test_export_uses_azure_translation_provider(
    monkeypatch: pytest.MonkeyPatch,
    sample_template_path: Path,
    sample_prompt: str,
    isolated_storage: Path,
) -> None:
    good_example_path = ROOT_DIR / "tests" / "fixtures" / "good_example.txt"
    bad_example_path = ROOT_DIR / "tests" / "fixtures" / "bad_example.txt"

    monkeypatch.setattr(rag_service, "index_examples", lambda *args, **kwargs: [])
    monkeypatch.setattr(translation_service.settings, "translation_provider", "azure")

    async def fake_generate_json(*, system_prompt: str, user_prompt: str, temperature: float = 0.2):
        if "assistant_message" in user_prompt:
            return {
                "assistant_message": "Captured the core scope.",
                "summary": "Compliance onboarding workflow for enterprise customers.",
                "section_updates": [
                    {
                        "section_id": "section-1",
                        "title": "Project Overview",
                        "content": "Build a compliance onboarding workflow for enterprise customers.",
                        "status": "complete",
                    },
                    {
                        "section_id": "section-2",
                        "title": "Functional Requirements",
                        "content": "Include identity verification, document review, approval audit trail, and reviewer notifications.",
                        "status": "complete",
                    },
                ],
            }
        raise AssertionError("The LLM provider should not be used for translation when Azure is active.")

    async def fake_translate_with_azure(*, session, language: str):
        assert language == "Spanish"
        return {
            "section-1::title": "Resumen del proyecto",
            "section-1": "Crear un flujo de incorporacion de cumplimiento para clientes empresariales.",
            "section-1::content": "Crear un flujo de incorporacion de cumplimiento para clientes empresariales.",
            "section-2::title": "Requisitos funcionales",
            "section-2": "Incluir verificacion de identidad, revision documental, trazabilidad de aprobacion y notificaciones.",
            "section-2::content": "Incluir verificacion de identidad, revision documental, trazabilidad de aprobacion y notificaciones.",
        }

    monkeypatch.setattr(llm_provider, "generate_json", fake_generate_json)
    monkeypatch.setattr(translation_service, "_translate_with_azure", fake_translate_with_azure)

    with TestClient(app) as client:
        result = run_batch_workflow(
            base_url="http://testserver",
            template_path=sample_template_path,
            good_example_paths=[good_example_path],
            bad_example_paths=[bad_example_path],
            message=sample_prompt,
            languages=["Spanish"],
            client=client,
        )

    export_payload = result["export"]
    archive_path = Path(export_payload["archive_path"])
    assert archive_path.exists()
    generated_document = Document(str(archive_path.parent / "template.spanish.docx"))
    generated_text = "\n".join(paragraph.text for paragraph in generated_document.paragraphs)
    assert "Resumen del proyecto" in generated_text
    assert "Crear un flujo de incorporacion de cumplimiento" in generated_text


def test_ingest_with_existing_document_is_saved_and_loaded_with_scenario(
    monkeypatch: pytest.MonkeyPatch,
    sample_template_path: Path,
    isolated_storage: Path,
    tmp_path: Path,
) -> None:
    monkeypatch.setattr(rag_service, "index_examples", lambda *args, **kwargs: [])

    existing_document_path = tmp_path / "existing-spec.docx"
    existing_document = Document()
    existing_document.add_heading("Project Overview", level=1)
    existing_document.add_paragraph("Existing onboarding workflow ready for further enhancement.")
    existing_document.save(str(existing_document_path))

    with TestClient(app) as client:
        with sample_template_path.open("rb") as template_handle, existing_document_path.open("rb") as existing_handle:
            ingest_response = client.post(
                "/ingest",
                files=[
                    (
                        "template",
                        (
                            sample_template_path.name,
                            template_handle.read(),
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        ),
                    ),
                    (
                        "existing_document",
                        (
                            existing_document_path.name,
                            existing_handle.read(),
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        ),
                    ),
                ],
            )
        ingest_response.raise_for_status()
        ingest_payload = ingest_response.json()
        assert ingest_payload["draft_state"]["sections"][0]["content"] == "Existing onboarding workflow ready for further enhancement."
        assert any(file["kind"] == "enhancement_document" for file in ingest_payload["loaded_files"])

        save_response = client.post(
            "/scenarios/save",
            json={
                "session_id": ingest_payload["session_id"],
                "scenario_id": "enhancement-scenario",
                "prompt": "",
                "target_languages": ["Spanish"],
                "output_file_name": "enhanced-output",
            },
        )
        save_response.raise_for_status()

        load_response = client.post("/scenarios/load", json={"scenario_id": "enhancement-scenario"})
        load_response.raise_for_status()
        load_payload = load_response.json()
        assert any(file["kind"] == "enhancement_document" for file in load_payload["loaded_files"])
        assert load_payload["draft_state"]["sections"][0]["content"] == "Existing onboarding workflow ready for further enhancement."


def test_custom_css_can_be_uploaded_and_cleared(isolated_storage: Path) -> None:
    css_text = ":root { --brand-accent: #123456; }\nbody { background-color: rgb(9, 17, 34) !important; }"

    with TestClient(app) as client:
        upload_response = client.post(
            "/config/custom-css",
            files={"custom_css": ("brand.css", css_text.encode("utf-8"), "text/css")},
        )
        upload_response.raise_for_status()
        upload_payload = upload_response.json()
        assert upload_payload["custom_css"]["enabled"] is True
        assert upload_payload["custom_css"]["file_name"] == "brand.css"

        get_response = client.get("/config/custom-css")
        get_response.raise_for_status()
        get_payload = get_response.json()
        assert get_payload["enabled"] is True
        assert get_payload["file_name"] == "brand.css"
        assert get_payload["css_text"] == css_text

        clear_response = client.delete("/config/custom-css")
        clear_response.raise_for_status()
        clear_payload = clear_response.json()
        assert clear_payload["custom_css"]["enabled"] is False

        final_response = client.get("/config/custom-css")
        final_response.raise_for_status()
        final_payload = final_response.json()
        assert final_payload["enabled"] is False
        assert final_payload["css_text"] is None