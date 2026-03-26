from __future__ import annotations

import io
import json
import sqlite3
import shutil
import sys
import zipfile
from pathlib import Path

from docx import Document
from pypdf import PdfReader
import pytest
from fastapi.testclient import TestClient

ROOT_DIR = Path(__file__).resolve().parents[1]
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

from app.main import app
from app.models.document_state import McpServerCatalogResponse, McpServerSummary
from app.services.docker_mcp_service import docker_mcp_service
from app.services.llm_provider import llm_provider
from app.services.rag_service import rag_service
from scripts.run_batch_workflow import run_batch_workflow


def test_save_and_load_scenario_persists_prompt_languages_and_draft(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
    docx_prompt: str,
    sample_docx_template_path: Path,
    sample_docx_good_example_path: Path,
    sample_docx_bad_example_path: Path,
    isolated_storage: Path,
) -> None:
    working_template_path = tmp_path / sample_docx_template_path.name
    shutil.copyfile(sample_docx_template_path, working_template_path)
    db_path = (isolated_storage / "scenarios" / "scenarios.db")

    monkeypatch.setattr(
        docker_mcp_service,
        "list_servers",
        lambda: McpServerCatalogResponse(
            available=True,
            servers=[McpServerSummary(name="fetch", description="Fetches URLs."), McpServerSummary(name="youtube_transcript")],
        ),
    )

    monkeypatch.setattr(rag_service, "index_examples", lambda *args, **kwargs: [])

    async def fake_generate_json(*, system_prompt: str, user_prompt: str, temperature: float = 0.2, mcp_servers=None):
        if '"assistant_message"' in user_prompt:
            return {
                "assistant_message": "Draft updated.",
                "summary": "Scenario draft summary.",
                "section_updates": [
                    {
                        "section_id": "section-1",
                        "title": "Project Overview",
                        "content": "Saved scenario overview.",
                        "status": "complete",
                    }
                ],
            }
        if "Target language: Spanish." in system_prompt:
            return {"sections": [{"section_id": "section-1", "title": "Resumen del proyecto", "content": "Resumen guardado."}]}
        raise AssertionError(f"Unexpected LLM prompt: {system_prompt}")

    monkeypatch.setattr(llm_provider, "generate_json", fake_generate_json)

    with TestClient(app) as client:
        result = run_batch_workflow(
            base_url="http://testserver",
            template_path=working_template_path,
            good_example_paths=[sample_docx_good_example_path],
            bad_example_paths=[sample_docx_bad_example_path],
            message=docx_prompt,
            languages=["Spanish"],
            output_file_name="customer-onboarding-spec",
            client=client,
        )

        session_id = result["ingest"]["session_id"]
        save_response = client.post(
            "/scenarios/save",
            json={
                "session_id": session_id,
                "scenario_id": "Nuno Scenario 01",
                "prompt": docx_prompt,
                "prompt_sequence": [docx_prompt, "Add a dedicated security section after the overview."],
                "mcp_servers": ["fetch", "youtube_transcript"],
                "target_languages": ["Spanish", "French"],
                "output_file_name": "customer-onboarding-spec",
            },
        )
        save_response.raise_for_status()

        list_response = client.get("/scenarios")
        list_response.raise_for_status()
        load_response = client.post("/scenarios/load", json={"scenario_id": "nuno-scenario-01"})
        load_response.raise_for_status()

    saved_payload = save_response.json()
    scenarios = list_response.json()
    loaded_payload = load_response.json()
    with sqlite3.connect(db_path) as connection:
        row = connection.execute(
            "SELECT output_file_name, target_languages_json, mcp_servers_json, prompt FROM scenarios WHERE scenario_id = ?",
            ("nuno-scenario-01",),
        ).fetchone()

    assert saved_payload["scenario_id"] == "nuno-scenario-01"
    assert saved_payload["output_file_name"] == "customer-onboarding-spec"
    assert any(scenario["scenario_id"] == "nuno-scenario-01" for scenario in scenarios)
    assert loaded_payload["scenario_id"] == "nuno-scenario-01"
    assert saved_payload["prompt_sequence"] == [docx_prompt, "Add a dedicated security section after the overview."]
    assert loaded_payload["prompt"] == f"{docx_prompt}\n\nAdd a dedicated security section after the overview."
    assert loaded_payload["prompt_sequence"] == [docx_prompt, "Add a dedicated security section after the overview."]
    assert loaded_payload["mcp_servers"] == ["fetch", "youtube_transcript"]
    assert loaded_payload["target_languages"] == ["Spanish", "French"]
    assert loaded_payload["output_file_name"] == "customer-onboarding-spec"
    assert loaded_payload["draft_state"]["sections"][0]["content"] == "Saved scenario overview."
    assert "Saved scenario overview." in loaded_payload["preview_markdown"]
    assert len(loaded_payload["loaded_files"]) == 3
    assert {item["kind"] for item in loaded_payload["loaded_files"]} == {"template", "good_example", "bad_example"}
    assert loaded_payload["loaded_files"][0]["download_path"].startswith(f"/sessions/{loaded_payload['session_id']}/files/")
    assert db_path.exists()
    assert row is not None
    assert row[0] == "customer-onboarding-spec"
    assert "Spanish" in row[1]
    assert "fetch" in row[2]
    assert json.loads(row[3]) == [docx_prompt, "Add a dedicated security section after the overview."]


def test_save_and_load_scenario_restores_generated_export_artifacts(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
    docx_prompt: str,
    sample_docx_template_path: Path,
    sample_docx_good_example_path: Path,
    sample_docx_bad_example_path: Path,
    isolated_storage: Path,
) -> None:
    working_template_path = tmp_path / sample_docx_template_path.name
    shutil.copyfile(sample_docx_template_path, working_template_path)

    monkeypatch.setattr(rag_service, "index_examples", lambda *args, **kwargs: [])

    async def fake_generate_json(*, system_prompt: str, user_prompt: str, temperature: float = 0.2, mcp_servers=None):
        if '"assistant_message"' in user_prompt:
            return {
                "assistant_message": "Draft updated.",
                "summary": "Scenario draft summary.",
                "section_updates": [
                    {
                        "section_id": "section-1",
                        "title": "Project Overview",
                        "content": "Saved scenario overview.",
                        "status": "complete",
                    }
                ],
            }
        if "Target language: Spanish." in system_prompt:
            return {"sections": [{"section_id": "section-1", "title": "Resumen del proyecto", "content": "Resumen guardado."}]}
        raise AssertionError(f"Unexpected LLM prompt: {system_prompt}")

    monkeypatch.setattr(llm_provider, "generate_json", fake_generate_json)

    with TestClient(app) as client:
        result = run_batch_workflow(
            base_url="http://testserver",
            template_path=working_template_path,
            good_example_paths=[sample_docx_good_example_path],
            bad_example_paths=[sample_docx_bad_example_path],
            message=docx_prompt,
            languages=["Spanish"],
            output_file_name="customer-onboarding-spec",
            client=client,
        )

        exported_session_id = result["ingest"]["session_id"]
        save_response = client.post(
            "/scenarios/save",
            json={
                "session_id": exported_session_id,
                "scenario_id": "Scenario With Export Files",
                "prompt": docx_prompt,
                "target_languages": ["Spanish"],
                "output_file_name": "customer-onboarding-spec",
            },
        )
        save_response.raise_for_status()

        load_response = client.post("/scenarios/load", json={"scenario_id": "scenario-with-export-files"})
        load_response.raise_for_status()
        loaded_session_id = load_response.json()["session_id"]

        list_exports_response = client.get(f"/export/{loaded_session_id}/files")
        list_exports_response.raise_for_status()

        archive_response = client.get(f"/export/{loaded_session_id}/download")
        archive_response.raise_for_status()

        export_files = list_exports_response.json()
        file_download_response = client.get(export_files[0]["download_path"])
        file_download_response.raise_for_status()

    scenario_generated_dir = isolated_storage / "scenarios" / "scenario-with-export-files" / "generated"
    restored_generated_dir = isolated_storage / "storage" / "generated" / loaded_session_id

    assert scenario_generated_dir.exists()
    assert (scenario_generated_dir / "customer-onboarding-spec.zip").exists()
    assert (scenario_generated_dir / "customer-onboarding-spec.spanish.docx").exists()

    assert restored_generated_dir.exists()
    assert (restored_generated_dir / "customer-onboarding-spec.zip").exists()
    assert (restored_generated_dir / "customer-onboarding-spec.spanish.docx").exists()

    assert len(export_files) == 1
    assert export_files[0]["language"] == "Spanish"
    assert export_files[0]["file_name"] == "customer-onboarding-spec.spanish.docx"

    with zipfile.ZipFile(io.BytesIO(archive_response.content)) as archive:
        assert "customer-onboarding-spec.spanish.docx" in set(archive.namelist())

    downloaded_document = Document(io.BytesIO(file_download_response.content))
    downloaded_text = "\n".join(paragraph.text for paragraph in downloaded_document.paragraphs)
    assert "Resumen guardado." in downloaded_text


def test_chat_replay_request_does_not_duplicate_saved_prompt_sequence(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
    docx_prompt: str,
    sample_docx_template_path: Path,
    sample_docx_good_example_path: Path,
    sample_docx_bad_example_path: Path,
) -> None:
    working_template_path = tmp_path / sample_docx_template_path.name
    shutil.copyfile(sample_docx_template_path, working_template_path)

    monkeypatch.setattr(rag_service, "index_examples", lambda *args, **kwargs: [])

    async def fake_generate_json(*, system_prompt: str, user_prompt: str, temperature: float = 0.2, mcp_servers=None):
        if '"assistant_message"' in user_prompt:
            return {
                "assistant_message": "Draft updated.",
                "summary": "Scenario draft summary.",
                "section_updates": [
                    {
                        "section_id": "section-1",
                        "title": "Project Overview",
                        "content": "Saved scenario overview.",
                        "status": "complete",
                    }
                ],
            }
        if "Target language: Spanish." in system_prompt:
            return {"sections": [{"section_id": "section-1", "title": "Resumen del proyecto", "content": "Resumen guardado."}]}
        raise AssertionError(f"Unexpected LLM prompt: {system_prompt}")

    monkeypatch.setattr(llm_provider, "generate_json", fake_generate_json)

    with TestClient(app) as client:
        result = run_batch_workflow(
            base_url="http://testserver",
            template_path=working_template_path,
            good_example_paths=[sample_docx_good_example_path],
            bad_example_paths=[sample_docx_bad_example_path],
            message=docx_prompt,
            languages=["Spanish"],
            output_file_name="customer-onboarding-spec",
            client=client,
        )

        session_id = result["ingest"]["session_id"]
        replay_response = client.post(
            "/chat",
            json={
                "session_id": session_id,
                "message": docx_prompt,
                "persist_prompt": False,
            },
        )
        replay_response.raise_for_status()

    replay_payload = replay_response.json()
    assert replay_payload["prompt_sequence"] == [docx_prompt]


@pytest.mark.parametrize("export_format", ["docx", "pdf"])
def test_full_scenario_lifecycle_generates_translated_artifacts_from_refined_draft(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
    sample_docx_template_path: Path,
    sample_docx_good_example_path: Path,
    sample_docx_bad_example_path: Path,
    sample_docx_enhancement_with_image_path: Path,
    isolated_storage: Path,
    export_format: str,
) -> None:
    working_template_path = tmp_path / sample_docx_template_path.name
    shutil.copyfile(sample_docx_template_path, working_template_path)

    monkeypatch.setattr(rag_service, "index_examples", lambda *args, **kwargs: [])

    refinement_payloads = {
        "Define the project overview for a regulated onboarding workflow.": {
            "assistant_message": "Captured the onboarding scope.",
            "summary": "Regulated onboarding workflow draft.",
            "section_updates": [
                {
                    "section_id": "section-1",
                    "title": "Project Overview",
                    "content": "Build a regulated onboarding workflow for enterprise customers across KYC, document review, and approval stages.",
                    "status": "complete",
                }
            ],
        },
        "Add functional and non-functional requirements for auditability and localization.": {
            "assistant_message": "Added the detailed requirements.",
            "summary": "Regulated onboarding workflow ready for export.",
            "section_updates": [
                {
                    "section_id": "section-2",
                    "title": "Functional Requirements",
                    "content": "Support identity verification, document review queues, dual approval routing, reviewer notifications, and exception handling.",
                    "status": "complete",
                },
                {
                    "section_id": "section-3",
                    "title": "Non-Functional Requirements",
                    "content": "Maintain a complete audit trail, enforce role-based access, retain evidence for seven years, and prepare all visible text for localization.",
                    "status": "complete",
                },
            ],
        },
    }

    async def fake_generate_json(*, system_prompt: str, user_prompt: str, temperature: float = 0.2, mcp_servers=None):
        for message, payload in refinement_payloads.items():
            if f"User message:\n{message}\n\n" in user_prompt:
                return payload
        if "Target language: Spanish." in system_prompt:
            return {
                "sections": [
                    {
                        "section_id": "section-1",
                        "title": "Resumen del proyecto",
                        "content": "Crear un flujo de onboarding regulado para clientes empresariales en KYC, revision documental y aprobacion.",
                    },
                    {
                        "section_id": "section-2",
                        "title": "Requisitos funcionales",
                        "content": "Admitir verificacion de identidad, colas de revision documental, aprobacion dual, notificaciones y gestion de excepciones.",
                    },
                    {
                        "section_id": "section-3",
                        "title": "Requisitos no funcionales",
                        "content": "Mantener una pista de auditoria completa, aplicar acceso por roles, conservar evidencia durante siete anos y preparar el texto para localizacion.",
                    },
                ]
            }
        if "Target language: French." in system_prompt:
            return {
                "sections": [
                    {
                        "section_id": "section-1",
                        "title": "Vue d'ensemble du projet",
                        "content": "Creer un flux d'onboarding reglemente pour les clients entreprise couvrant KYC, revue documentaire et approbation.",
                    },
                    {
                        "section_id": "section-2",
                        "title": "Exigences fonctionnelles",
                        "content": "Prendre en charge la verification d'identite, les files de revue documentaire, l'approbation double, les notifications et la gestion des exceptions.",
                    },
                    {
                        "section_id": "section-3",
                        "title": "Exigences non fonctionnelles",
                        "content": "Maintenir une piste d'audit complete, appliquer l'acces par roles, conserver les preuves pendant sept ans et preparer le texte pour la localisation.",
                    },
                ]
            }
        raise AssertionError(f"Unexpected LLM prompt: {system_prompt}")

    monkeypatch.setattr(llm_provider, "generate_json", fake_generate_json)

    with TestClient(app) as client:
        with (
            working_template_path.open("rb") as template_handle,
            sample_docx_enhancement_with_image_path.open("rb") as enhancement_handle,
            sample_docx_good_example_path.open("rb") as good_handle,
            sample_docx_bad_example_path.open("rb") as bad_handle,
        ):
            ingest_response = client.post(
                "/ingest",
                files=[
                    (
                        "template",
                        (
                            working_template_path.name,
                            template_handle.read(),
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        ),
                    ),
                    (
                        "existing_document",
                        (
                            sample_docx_enhancement_with_image_path.name,
                            enhancement_handle.read(),
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        ),
                    ),
                    (
                        "good_examples",
                        (
                            sample_docx_good_example_path.name,
                            good_handle.read(),
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        ),
                    ),
                    (
                        "bad_examples",
                        (
                            sample_docx_bad_example_path.name,
                            bad_handle.read(),
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        ),
                    ),
                ],
            )
        ingest_response.raise_for_status()
        ingest_payload = ingest_response.json()
        session_id = ingest_payload["session_id"]

        first_chat_response = client.post(
            "/chat",
            json={
                "session_id": session_id,
                "message": "Define the project overview for a regulated onboarding workflow.",
            },
        )
        first_chat_response.raise_for_status()

        second_chat_response = client.post(
            "/chat",
            json={
                "session_id": session_id,
                "message": "Add functional and non-functional requirements for auditability and localization.",
            },
        )
        second_chat_response.raise_for_status()
        second_chat_payload = second_chat_response.json()

        save_response = client.post(
            "/scenarios/save",
            json={
                "session_id": session_id,
                "scenario_id": f"end-to-end-{export_format}",
                "prompt_sequence": [
                    "Define the project overview for a regulated onboarding workflow.",
                    "Add functional and non-functional requirements for auditability and localization.",
                ],
                "target_languages": ["Spanish", "French"],
                "output_file_name": f"end-to-end-{export_format}-spec",
            },
        )
        save_response.raise_for_status()

        load_response = client.post("/scenarios/load", json={"scenario_id": f"end-to-end-{export_format}"})
        load_response.raise_for_status()
        loaded_payload = load_response.json()

        export_response = client.post(
            "/export",
            json={
                "session_id": loaded_payload["session_id"],
                "target_languages": ["Spanish", "French"],
                "output_file_name": f"end-to-end-{export_format}-spec",
                "export_format": export_format,
            },
        )
        export_response.raise_for_status()
        export_payload = export_response.json()

        list_exports_response = client.get(f"/export/{loaded_payload['session_id']}/files")
        list_exports_response.raise_for_status()
        listed_files = list_exports_response.json()

    assert "regulated onboarding workflow" in second_chat_payload["preview_markdown"].lower()
    assert "audit trail" in second_chat_payload["preview_markdown"].lower()
    assert loaded_payload["prompt_sequence"] == [
        "Define the project overview for a regulated onboarding workflow.",
        "Add functional and non-functional requirements for auditability and localization.",
    ]

    assert export_payload["export_format"] == export_format
    assert len(export_payload["generated_documents"]) == 2
    assert {document["language"] for document in export_payload["generated_documents"]} == {"Spanish", "French"}
    assert {document["format"] for document in export_payload["generated_documents"]} == {export_format}
    assert len(listed_files) == 2

    generated_paths = [Path(file_path) for file_path in export_payload["generated_files"]]
    assert all(path.exists() for path in generated_paths)

    archive_path = Path(export_payload["archive_path"])
    assert archive_path.exists()
    with zipfile.ZipFile(archive_path) as archive:
        archive_names = set(archive.namelist())
        assert f"end-to-end-{export_format}-spec.spanish.{export_format}" in archive_names
        assert f"end-to-end-{export_format}-spec.french.{export_format}" in archive_names

    if export_format == "docx":
        spanish_path = next(path for path in generated_paths if path.name.endswith(".spanish.docx"))
        french_path = next(path for path in generated_paths if path.name.endswith(".french.docx"))
        spanish_text = "\n".join(paragraph.text for paragraph in Document(str(spanish_path)).paragraphs if paragraph.text.strip())
        french_text = "\n".join(paragraph.text for paragraph in Document(str(french_path)).paragraphs if paragraph.text.strip())
        assert "Resumen del proyecto" in spanish_text
        assert "pista de auditoria" in spanish_text
        assert "Vue d'ensemble du projet" in french_text
        assert "piste d'audit complete" in french_text
    else:
        assert all(path.read_bytes().startswith(b"%PDF") for path in generated_paths)
        spanish_path = next(path for path in generated_paths if path.name.endswith(".spanish.pdf"))
        french_path = next(path for path in generated_paths if path.name.endswith(".french.pdf"))
        spanish_text = "\n".join((page.extract_text() or "") for page in PdfReader(str(spanish_path)).pages)
        french_text = "\n".join((page.extract_text() or "") for page in PdfReader(str(french_path)).pages)
        assert "Resumen del proyecto" in spanish_text
        assert "pista de auditoria completa" in spanish_text
        assert "Vue d'ensemble du projet" in french_text
        assert "piste d'audit complete" in french_text