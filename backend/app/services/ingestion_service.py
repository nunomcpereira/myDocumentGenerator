from __future__ import annotations

from dataclasses import dataclass
from datetime import UTC, datetime
from pathlib import Path
import re
from uuid import uuid4

from fastapi import UploadFile
from pypdf import PdfReader

from app.core.config import get_settings
from app.core.errors import UnsupportedTemplateError
from app.models.document_state import DraftSectionState, DocumentDraftState, SessionContext, TemplateSection, TemplateStructure
from app.services.markitdown_service import markitdown_service
from app.services.rag_service import rag_service

try:
    from docx import Document as DocxDocument
except ImportError:  # pragma: no cover - optional dependency path
    DocxDocument = None


@dataclass
class ImportedSection:
    title: str
    content: str
    image_paths: list[Path]
    level: int = 1


class IngestionService:
    def __init__(self) -> None:
        self.settings = get_settings()

    async def initialize_session(
        self,
        *,
        template_file: UploadFile,
        enhancement_document_file: UploadFile | None,
        good_examples: list[UploadFile],
        bad_examples: list[UploadFile],
    ) -> SessionContext:
        session_id = str(uuid4())
        session_dir = self.settings.upload_root / session_id
        session_dir.mkdir(parents=True, exist_ok=True)

        template_path = await self._persist_file(session_dir, template_file)
        template_structure = self._parse_template(template_path)
        enhancement_document_path = await self._persist_file(session_dir, enhancement_document_file) if enhancement_document_file else None

        saved_good_examples = [await self._persist_file(session_dir, example) for example in good_examples]
        saved_bad_examples = [await self._persist_file(session_dir, example) for example in bad_examples]
        warnings = self.index_session_examples(
            session_id,
            saved_good_examples,
            saved_bad_examples,
            enhancement_document_path=enhancement_document_path,
        )

        draft_state, draft_warnings = self._build_initial_draft_state(
            session_id=session_id,
            template_structure=template_structure,
            enhancement_document_path=enhancement_document_path,
        )
        warnings.extend(draft_warnings)
        enhancement_image_paths, enhancement_section_image_paths, image_warnings = self.extract_preview_assets(
            session_id=session_id,
            template_path=template_path,
            template_structure=template_structure,
            enhancement_document_path=enhancement_document_path,
        )
        warnings.extend(image_warnings)
        source_preview_markdown = self.build_source_preview_markdown(enhancement_document_path or template_path)

        return SessionContext(
            session_id=session_id,
            template_path=template_path,
            template_structure=template_structure,
            original_template_structure=template_structure.model_copy(deep=True),
            source_preview_markdown=source_preview_markdown,
            draft_state=draft_state,
            enhancement_document_path=enhancement_document_path,
            enhancement_image_paths=enhancement_image_paths,
            enhancement_section_image_paths=enhancement_section_image_paths,
            good_example_paths=saved_good_examples,
            bad_example_paths=saved_bad_examples,
            output_file_name=template_path.stem,
            warnings=warnings,
        )

    async def _persist_file(self, directory: Path, upload: UploadFile) -> Path:
        file_path = directory / upload.filename
        content = await upload.read()
        file_path.write_bytes(content)
        await upload.close()
        return file_path

    def index_session_examples(
        self,
        session_id: str,
        good_example_paths: list[Path],
        bad_example_paths: list[Path],
        *,
        enhancement_document_path: Path | None = None,
    ) -> list[str]:
        good_examples = [(path.name, self._read_example_text(path)) for path in good_example_paths]
        if enhancement_document_path is not None:
            good_examples.append((enhancement_document_path.name, self._read_example_text(enhancement_document_path)))
        bad_examples = [(path.name, self._read_example_text(path)) for path in bad_example_paths]
        return rag_service.index_examples(session_id, good_examples, bad_examples)

    def extract_preview_assets(
        self,
        *,
        session_id: str,
        template_path: Path,
        template_structure: TemplateStructure,
        enhancement_document_path: Path | None,
    ) -> tuple[list[Path], dict[str, list[Path]], list[str]]:
        source_path = enhancement_document_path or template_path
        suffix = source_path.suffix.lower()
        if suffix == ".docx" and DocxDocument is None:
            return [], {}, []
        if suffix not in {".docx", ".pdf"}:
            return [], {}, []

        image_output_dir = self.settings.upload_root / session_id / "preview_images"
        imported_sections = self._extract_existing_sections(source_path, image_output_dir=image_output_dir)
        warning_source = "enhancement document" if enhancement_document_path else "template"
        return self._build_section_image_mapping(template_structure, imported_sections, warning_source=warning_source)

    def build_source_preview_markdown(self, source_path: Path | None) -> str | None:
        if source_path is None:
            return None
        try:
            return markitdown_service.convert_file(source_path).markdown or None
        except Exception:
            return None

    def _build_initial_draft_state(
        self,
        *,
        session_id: str,
        template_structure: TemplateStructure,
        enhancement_document_path: Path | None,
    ) -> tuple[DocumentDraftState, list[str]]:
        draft_state = DocumentDraftState(
            session_id=session_id,
            sections=[
                DraftSectionState(section_id=section.id, title=section.title)
                for section in template_structure.sections
            ],
            updated_at=datetime.now(UTC),
        )

        if enhancement_document_path is None:
            return draft_state, []

        warnings = self._hydrate_draft_from_existing_document(draft_state, template_structure, enhancement_document_path)
        draft_state.updated_at = datetime.now(UTC)
        return draft_state, warnings

    def _hydrate_draft_from_existing_document(
        self,
        draft_state: DocumentDraftState,
        template_structure: TemplateStructure,
        existing_document_path: Path,
    ) -> list[str]:
        warnings: list[str] = []
        template_sections = {self._normalize_title(section.title): section for section in template_structure.sections}
        draft_sections = {section.section_id: section for section in draft_state.sections}
        imported_sections = self._extract_existing_sections(existing_document_path)
        matched = 0

        for imported_section in imported_sections:
            template_section = template_sections.get(self._normalize_title(imported_section.title))
            if template_section is None:
                continue
            draft_section = draft_sections.get(template_section.id)
            if draft_section is None:
                continue
            draft_section.title = imported_section.title
            if not imported_section.content.strip():
                continue
            draft_section.content = imported_section.content.strip()
            draft_section.status = "complete"
            draft_section.last_updated_at = datetime.now(UTC)
            matched += 1

        if matched == 0:
            imported_text = self._read_example_text(existing_document_path).strip()
            if imported_text and draft_state.sections:
                draft_state.sections[0].content = imported_text
                draft_state.sections[0].status = "in_progress"
                draft_state.sections[0].last_updated_at = datetime.now(UTC)
                warnings.append(
                    "The enhancement document could not be aligned to the template headings, so its content was placed in the first section for refinement."
                )
        else:
            draft_state.summary = f"Imported content from {existing_document_path.name} into {matched} template section(s)."

        return warnings

    def _extract_existing_sections(self, file_path: Path, image_output_dir: Path | None = None) -> list[ImportedSection]:
        suffix = file_path.suffix.lower()
        if suffix == ".docx" and DocxDocument is not None:
            from docx.text.paragraph import Paragraph as _DocxParagraph
            from docx.table import Table as _DocxTable

            document = DocxDocument(str(file_path))
            sections: list[ImportedSection] = []
            current_title: str | None = None
            current_lines: list[str] = []
            current_images: list[Path] = []
            image_index = 0

            def flush_current_section() -> None:
                if current_title and (current_lines or current_images):
                    sections.append(
                        ImportedSection(
                            title=current_title,
                            content="\n".join(current_lines).strip(),
                            image_paths=list(current_images),
                        )
                    )

            for child in document.element.body:
                local_tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

                if local_tag == "tbl":
                    table_md = self._docx_table_to_markdown(_DocxTable(child, document))
                    if table_md:
                        if current_title is None:
                            current_title = "Imported content"
                        current_lines.append(table_md)
                    continue

                if local_tag != "p":
                    continue

                paragraph = _DocxParagraph(child, document)
                text = paragraph.text.strip()
                style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
                paragraph_images, image_index = self._extract_docx_paragraph_images(
                    document,
                    paragraph,
                    image_output_dir,
                    image_index,
                )

                if style_name.startswith("heading") and text:
                    flush_current_section()
                    current_title = text
                    current_lines = []
                    current_images = []
                    continue

                if not text and not paragraph_images:
                    continue

                if current_title is None:
                    current_title = "Imported content"
                if text:
                    current_lines.append(text)
                if paragraph_images:
                    current_images.extend(paragraph_images)

            flush_current_section()
            return sections

        if suffix == ".pdf":
            try:
                converted = markitdown_service.convert_file(file_path)
            except Exception:
                return []
            sections = self._split_markdown_into_sections(
                converted.markdown,
                default_title=converted.title or self._default_imported_title(file_path),
            )
            if image_output_dir is not None and sections:
                reader = PdfReader(str(file_path))
                image_paths = self._extract_pdf_page_images(reader, image_output_dir)
                if image_paths:
                    sections[0].image_paths.extend(image_paths)
            return sections

        try:
            converted = markitdown_service.convert_file(file_path)
        except Exception:
            return []
        return self._split_markdown_into_sections(
            converted.markdown,
            default_title=converted.title or self._default_imported_title(file_path),
        )

    def _build_section_image_mapping(
        self,
        template_structure: TemplateStructure,
        imported_sections: list[ImportedSection],
        *,
        warning_source: str,
    ) -> tuple[list[Path], dict[str, list[Path]], list[str]]:
        image_paths_by_section: dict[str, list[Path]] = {}
        warnings: list[str] = []
        template_sections = {self._normalize_title(section.title): section for section in template_structure.sections}
        unmatched_images: list[Path] = []

        for imported_section in imported_sections:
            if not imported_section.image_paths:
                continue
            template_section = template_sections.get(self._normalize_title(imported_section.title))
            if template_section is None:
                unmatched_images.extend(imported_section.image_paths)
                continue
            image_paths_by_section.setdefault(template_section.id, []).extend(imported_section.image_paths)

        if unmatched_images and template_structure.sections:
            first_section_id = template_structure.sections[0].id
            image_paths_by_section.setdefault(first_section_id, []).extend(unmatched_images)
            warnings.append(
                f"Images from the {warning_source} could not be aligned to template headings, so they are shown in the first section preview."
            )

        all_image_paths: list[Path] = []
        for paths in image_paths_by_section.values():
            all_image_paths.extend(paths)
        return all_image_paths, image_paths_by_section, warnings

    @staticmethod
    def _extract_docx_paragraph_images(document: DocxDocument, paragraph, image_output_dir: Path | None, image_index: int) -> tuple[list[Path], int]:
        if image_output_dir is None:
            return [], image_index

        image_output_dir.mkdir(parents=True, exist_ok=True)
        image_paths: list[Path] = []
        for blip in paragraph._element.xpath(".//*[local-name()='blip']"):
            rel_id = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            if not rel_id:
                continue
            image_part = document.part.related_parts.get(rel_id)
            if image_part is None:
                continue
            image_index += 1
            suffix = Path(str(image_part.partname)).suffix or ".png"
            image_path = image_output_dir / f"preview-image-{image_index}{suffix}"
            image_path.write_bytes(image_part.blob)
            image_paths.append(image_path)
        return image_paths, image_index

    @staticmethod
    def _normalize_title(value: str) -> str:
        value_without_numbering = IngestionService._strip_leading_heading_number(value)
        return " ".join(re.sub(r"[_\W]+", " ", value_without_numbering.casefold(), flags=re.UNICODE).split())

    def _parse_template(self, template_path: Path) -> TemplateStructure:
        suffix = template_path.suffix.lower()
        if suffix == ".docx":
            return self._parse_docx_template(template_path)
        return self._parse_markitdown_template(template_path)

    def _parse_markitdown_template(self, template_path: Path) -> TemplateStructure:
        try:
            converted = markitdown_service.convert_file(template_path)
        except Exception as exc:
            raise UnsupportedTemplateError(
                f"{template_path.name} could not be converted to markdown. Upload a file supported by MarkItDown or use a .docx/.pdf template."
            ) from exc

        sections = self._split_markdown_into_sections(
            converted.markdown,
            default_title=converted.title or self._default_imported_title(template_path),
        )
        if not sections:
            sections = [
                ImportedSection(
                    title=converted.title or self._default_imported_title(template_path),
                    content=converted.markdown,
                    image_paths=[],
                )
            ]

        file_type = template_path.suffix.lower().lstrip(".") or "generic"
        outline = [section.title for section in sections]
        return TemplateStructure(
            file_name=template_path.name,
            file_type=file_type,
            sections=[
                TemplateSection(
                    id=f"section-{index}",
                    title=section.title,
                    level=section.level,
                    source_excerpt=(section.content or section.title)[:1200],
                )
                for index, section in enumerate(sections, start=1)
            ],
            extracted_outline=outline,
        )

    def _parse_docx_template(self, template_path: Path) -> TemplateStructure:
        if DocxDocument is None:
            raise UnsupportedTemplateError("python-docx is not installed, so .docx templates cannot be parsed.")

        document = DocxDocument(str(template_path))
        sections: list[TemplateSection] = []
        current_section: TemplateSection | None = None
        extracted_outline: list[str] = []

        for index, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
            if style_name.startswith("heading"):
                level = self._extract_heading_level(style_name)
                section_id = f"section-{len(sections) + 1}"
                current_section = TemplateSection(
                    id=section_id,
                    title=text,
                    level=level,
                    source_excerpt=text,
                    heading_paragraph_index=index,
                )
                sections.append(current_section)
                extracted_outline.append(text)
                continue

            if current_section is None:
                section_id = f"section-{len(sections) + 1}"
                current_section = TemplateSection(
                    id=section_id,
                    title=f"Section {len(sections) + 1}",
                    level=1,
                    source_excerpt=text,
                )
                sections.append(current_section)
                extracted_outline.append(current_section.title)
            current_section.content_paragraph_indices.append(index)
            excerpt = current_section.source_excerpt or ""
            current_section.source_excerpt = f"{excerpt}\n{text}".strip()

        return TemplateStructure(
            file_name=template_path.name,
            file_type="docx",
            sections=sections,
            extracted_outline=extracted_outline,
        )

    @staticmethod
    def _docx_table_to_markdown(table) -> str:
        """Convert a python-docx Table object to a GitHub-flavored markdown table string."""
        try:
            if not table.rows:
                return ""
            rows_md: list[str] = []
            for row in table.rows:
                cells = [" ".join(cell.text.split()).replace("|", "\\|") for cell in row.cells]
                rows_md.append("| " + " | ".join(cells) + " |")
            if not rows_md:
                return ""
            col_count = len(table.rows[0].cells)
            separator = "| " + " | ".join(["---"] * col_count) + " |"
            return "\n".join([rows_md[0], separator] + rows_md[1:])
        except Exception:
            return ""

    @staticmethod
    def _extract_pdf_page_images(reader, image_output_dir: Path) -> list[Path]:
        """Extract all embedded images from a pypdf PdfReader and save them to disk."""
        image_output_dir.mkdir(parents=True, exist_ok=True)
        image_paths: list[Path] = []
        image_index = 0
        for page in reader.pages:
            try:
                for img in page.images:
                    image_index += 1
                    suffix = Path(img.name).suffix if img.name else ".png"
                    if not suffix:
                        suffix = ".png"
                    image_path = image_output_dir / f"pdf-image-{image_index}{suffix}"
                    image_path.write_bytes(img.data)
                    image_paths.append(image_path)
            except Exception:
                continue
        return image_paths

    @staticmethod
    def _strip_leading_heading_number(value: str) -> str:
        return re.sub(r"^\d+(?:\.\d+){0,4}(?:[.)])?\s+", "", value).strip()

    def _read_example_text(self, file_path: Path) -> str:
        try:
            converted = markitdown_service.convert_file(file_path)
            if converted.markdown:
                return converted.markdown
        except Exception:
            pass

        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            reader = PdfReader(str(file_path))
            return "\n".join((page.extract_text() or "") for page in reader.pages).strip()
        if suffix == ".docx" and DocxDocument is not None:
            document = DocxDocument(str(file_path))
            return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()).strip()
        return file_path.read_text(encoding="utf-8", errors="ignore")

    @staticmethod
    def _default_imported_title(file_path: Path) -> str:
        return file_path.stem.replace("_", " ").replace("-", " ").strip() or "Imported content"

    @staticmethod
    def _split_markdown_into_sections(markdown: str, *, default_title: str) -> list[ImportedSection]:
        if not markdown.strip():
            return []

        sections: list[ImportedSection] = []
        current_title: str | None = None
        current_level = 1
        current_lines: list[str] = []

        def flush_current_section() -> None:
            if current_title is None:
                return
            sections.append(
                ImportedSection(
                    title=current_title,
                    content="\n".join(current_lines).strip(),
                    image_paths=[],
                    level=current_level,
                )
            )

        for raw_line in markdown.splitlines():
            line = raw_line.rstrip()
            heading_match = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
            if heading_match is not None:
                flush_current_section()
                current_title = heading_match.group(2).strip()
                current_level = len(heading_match.group(1))
                current_lines = []
                continue

            if current_title is None:
                current_title = default_title
                current_level = 1
            current_lines.append(line)

        flush_current_section()

        normalized_sections = [
            ImportedSection(
                title=section.title,
                content=section.content.strip(),
                image_paths=section.image_paths,
                level=section.level,
            )
            for section in sections
            if section.title.strip() or section.content.strip()
        ]
        return normalized_sections

    @staticmethod
    def _extract_heading_level(style_name: str) -> int:
        for token in style_name.split():
            if token.isdigit():
                return int(token)
        digits = "".join(character for character in style_name if character.isdigit())
        return int(digits) if digits else 1


ingestion_service = IngestionService()
