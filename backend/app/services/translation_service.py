from __future__ import annotations

from html import unescape
import re
import zipfile
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape as xml_escape

import httpx

from app.core.config import get_settings
from app.core.errors import ExportError
from app.core.errors import TranslationProviderConfigurationError, TranslationProviderError
from app.models.document_state import CustomCssConfiguration, ExampleSnippet, SessionContext, TemplateSection, TemplateStructure, TranslationConfigurationResponse, TranslationProviderId, TranslationProviderOption
from app.services.llm_provider import llm_provider
from app.services.scenario_service import scenario_service

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches as DocxInches
except ImportError:  # pragma: no cover - optional dependency path
    DocxDocument = None
    DocxInches = None

try:
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Image as RLImage, Paragraph, SimpleDocTemplate, Spacer
except ImportError:  # pragma: no cover - optional dependency path
    RLImage = None
    Paragraph = None
    ParagraphStyle = None
    SimpleDocTemplate = None
    Spacer = None
    getSampleStyleSheet = None


LANGUAGE_CODES = {
    "arabic": "ar",
    "chinese": "zh-Hans",
    "dutch": "nl",
    "english": "en",
    "french": "fr",
    "german": "de",
    "italian": "it",
    "japanese": "ja",
    "korean": "ko",
    "portuguese": "pt",
    "spanish": "es",
}

PROVIDER_METADATA: dict[TranslationProviderId, dict[str, Any]] = {
    "llm": {
        "label": "Current LLM",
        "description": "Uses the existing JSON-based translation prompt against the configured chat model.",
        "required_env": ["TRANSLATION_PROVIDER", "LLM_BASE_URL", "LLM_API_KEY", "LLM_MODEL"],
    },
    "azure": {
        "label": "Azure Translator",
        "description": "Calls the Azure AI Translator Text API directly for section title and body translation.",
        "required_env": ["TRANSLATION_PROVIDER", "AZURE_TRANSLATOR_ENDPOINT", "AZURE_TRANSLATOR_KEY", "AZURE_TRANSLATOR_REGION"],
    },
    "google": {
        "label": "Google Translate",
        "description": "Calls the Google Cloud Translation REST API directly for section title and body translation.",
        "required_env": ["TRANSLATION_PROVIDER", "GOOGLE_TRANSLATE_API_KEY"],
    },
}


class TranslationService:
    def __init__(self) -> None:
        self.settings = get_settings()

    def describe_configuration(self) -> TranslationConfigurationResponse:
        active_provider = self.get_active_provider()
        _, css_file_name, css_updated_at = scenario_service.get_custom_css()
        return TranslationConfigurationResponse(
            active_provider=active_provider,
            options=[
                TranslationProviderOption(
                    id=provider_id,
                    label=metadata["label"],
                    description=metadata["description"],
                    configured=self._is_provider_configured(provider_id),
                    required_env=metadata["required_env"],
                )
                for provider_id, metadata in PROVIDER_METADATA.items()
            ],
            custom_css=CustomCssConfiguration(enabled=bool(css_file_name), file_name=css_file_name, updated_at=css_updated_at),
        )

    def get_active_provider(self) -> TranslationProviderId:
        raw_provider = (self.settings.translation_provider or "llm").strip().lower()
        if raw_provider in PROVIDER_METADATA:
            return raw_provider  # type: ignore[return-value]
        raise TranslationProviderConfigurationError(
            "Unsupported translation provider configured. Use TRANSLATION_PROVIDER=llm, azure, or google."
        )

    async def translate_sections(
        self,
        *,
        session: SessionContext,
        language: str,
        good_examples: list[ExampleSnippet],
    ) -> dict[str, str]:
        provider = self.get_active_provider()
        if provider == "llm":
            return await self._translate_with_llm(session=session, language=language, good_examples=good_examples)
        if provider == "azure":
            return await self._translate_with_azure(session=session, language=language)
        return await self._translate_with_google(session=session, language=language)

    def inject_translations(
        self,
        *,
        session: SessionContext,
        translations: dict[str, dict[str, str]],
        output_directory: Path,
        output_file_name: str | None = None,
        export_format: str = "docx",
    ) -> tuple[Path, list[Path], list[str]]:
        self.validate_export_request(export_format)

        output_directory.mkdir(parents=True, exist_ok=True)
        self._clear_generated_outputs(output_directory)
        generated_files: list[Path] = []
        warnings: list[str] = []
        base_name = (output_file_name or session.output_file_name or session.template_path.stem).strip() or session.template_path.stem

        for language, translated_sections in translations.items():
            sections = self._build_export_sections(session=session, translated_sections=translated_sections)
            extension = "pdf" if export_format == "pdf" else "docx"
            target_path = output_directory / f"{base_name}.{language.lower()}.{extension}"
            if export_format == "pdf":
                self._write_pdf_export(target_path=target_path, sections=sections)
            else:
                self._write_docx_export(target_path=target_path, sections=sections)
            generated_files.append(target_path)

        archive_path = output_directory / f"{base_name}.zip"
        with zipfile.ZipFile(archive_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
            for file_path in generated_files:
                archive.write(file_path, arcname=file_path.name)
        return archive_path, generated_files, warnings

    def validate_export_request(self, export_format: str) -> None:
        if export_format not in {"docx", "pdf"}:
            raise ExportError("Unsupported export format. Use docx or pdf.")
        if export_format == "docx" and DocxDocument is None:
            raise ExportError("python-docx is not installed, so DOCX export is unavailable.")
        if export_format == "pdf" and (SimpleDocTemplate is None or Paragraph is None or Spacer is None or getSampleStyleSheet is None or ParagraphStyle is None):
            raise ExportError("PDF export requires ReportLab in the backend environment.")

    @staticmethod
    def _clear_generated_outputs(output_directory: Path) -> None:
        for pattern in ("*.docx", "*.pdf", "*.zip"):
            for file_path in output_directory.glob(pattern):
                if file_path.is_file():
                    file_path.unlink()

    def _build_export_sections(self, *, session: SessionContext, translated_sections: dict[str, str]) -> list[dict[str, Any]]:
        template_sections = {section.id: section for section in session.template_structure.sections}
        export_sections: list[dict[str, Any]] = []
        for draft_section in session.draft_state.sections:
            template_section = template_sections.get(draft_section.section_id)
            title = translated_sections.get(f"{draft_section.section_id}::title") or draft_section.title
            content = translated_sections.get(f"{draft_section.section_id}::content") or translated_sections.get(draft_section.section_id) or draft_section.content
            image_paths = session.enhancement_section_image_paths.get(draft_section.section_id, [])
            export_sections.append(
                {
                    "title": title,
                    "content": content,
                    "level": max(1, min((template_section.level if template_section else 1), 9)),
                    "image_paths": [p for p in image_paths if p.exists()],
                }
            )
        return export_sections

    def _write_docx_export(self, *, target_path: Path, sections: list[dict[str, Any]]) -> None:
        document = DocxDocument()
        for section in sections:
            document.add_heading(section["title"], level=section["level"])
            self._append_docx_markdownish_content(document, section["content"])
            for image_path in section.get("image_paths", []):
                try:
                    kwargs = {"width": DocxInches(5)} if DocxInches is not None else {}
                    document.add_picture(str(image_path), **kwargs)
                except Exception:
                    pass
        document.save(str(target_path))

    def _append_docx_markdownish_content(self, document: DocxDocument, content: str) -> None:
        blocks = [block.strip() for block in re.split(r"\n\s*\n", content or "") if block.strip()]
        if not blocks:
            document.add_paragraph("")
            return
        for block in blocks:
            lines = [line.rstrip() for line in block.splitlines() if line.strip()]
            if not lines:
                continue
            # Markdown table: first line is |...|, second is |---|
            if len(lines) >= 2 and re.match(r"^\|.+\|$", lines[0].strip()) and re.match(r"^\|[-| :]+\|$", lines[1].strip()):
                self._append_docx_markdown_table(document, lines)
                continue
            if all(line.lstrip().startswith(("- ", "* ")) for line in lines):
                for line in lines:
                    para = document.add_paragraph(style="List Bullet")
                    self._add_inline_markdown_runs(para, line.lstrip()[2:].strip())
                continue
            if all(re.match(r"\d+\.\s+", line.lstrip()) for line in lines):
                for line in lines:
                    para = document.add_paragraph(style="List Number")
                    self._add_inline_markdown_runs(para, re.sub(r"^\d+\.\s+", "", line.lstrip()))
                continue
            para = document.add_paragraph()
            self._add_inline_markdown_runs(para, "\n".join(lines))

    @staticmethod
    def _add_inline_markdown_runs(paragraph: Any, text: str) -> None:
        """Add runs to a paragraph, parsing **bold** and *italic* inline markdown."""
        token_pattern = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*|_(.+?)_")
        last_end = 0
        for match in token_pattern.finditer(text):
            start, end = match.span()
            if start > last_end:
                paragraph.add_run(text[last_end:start])
            if match.group(1) is not None:
                paragraph.add_run(match.group(1)).bold = True
            else:
                paragraph.add_run(match.group(2) or match.group(3)).italic = True
            last_end = end
        if last_end < len(text):
            paragraph.add_run(text[last_end:])

    @staticmethod
    def _append_docx_markdown_table(document: DocxDocument, lines: list[str]) -> None:
        """Render a markdown table into a python-docx table."""
        try:
            header_cells = [cell.strip() for cell in lines[0].strip("|").split("|")]
            data_rows = [
                [cell.strip() for cell in line.strip("|").split("|")]
                for line in lines[2:]
                if line.strip() and not re.match(r"^\|[-| :]+\|$", line.strip())
            ]
            col_count = len(header_cells)
            if col_count == 0:
                return
            table = document.add_table(rows=1 + len(data_rows), cols=col_count)
            table.style = "Table Grid"
            for i, cell_text in enumerate(header_cells[:col_count]):
                cell = table.rows[0].cells[i]
                cell.text = cell_text
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for row_idx, data_row in enumerate(data_rows):
                for col_idx, cell_text in enumerate(data_row[:col_count]):
                    table.rows[row_idx + 1].cells[col_idx].text = cell_text
        except Exception:
            pass

    def _write_pdf_export(self, *, target_path: Path, sections: list[dict[str, Any]]) -> None:
        stylesheet = getSampleStyleSheet()
        body_style = stylesheet["BodyText"]
        heading_styles = {
            1: stylesheet["Heading1"],
            2: stylesheet["Heading2"],
            3: stylesheet["Heading3"],
        }
        bullet_style = ParagraphStyle(
            "ExportBullet",
            parent=body_style,
            leftIndent=18,
            firstLineIndent=-10,
            spaceBefore=2,
            spaceAfter=2,
        )
        number_style = ParagraphStyle(
            "ExportNumber",
            parent=body_style,
            leftIndent=18,
            firstLineIndent=-10,
            spaceBefore=2,
            spaceAfter=2,
        )

        story: list[Any] = []
        for section in sections:
            story.append(Paragraph(xml_escape(section["title"]), heading_styles.get(section["level"], stylesheet["Heading3"])))
            story.extend(self._build_pdf_story_blocks(section["content"], body_style, bullet_style, number_style))
            for image_path in section.get("image_paths", []):
                try:
                    img = RLImage(str(image_path))
                    max_width = 4.5 * inch
                    if img.drawWidth > max_width:
                        scale = max_width / img.drawWidth
                        img.drawWidth = max_width
                        img.drawHeight = img.drawHeight * scale
                    story.append(img)
                    story.append(Spacer(1, 0.1 * inch))
                except Exception:
                    pass
            story.append(Spacer(1, 0.18 * inch))

        document = SimpleDocTemplate(
            str(target_path),
            leftMargin=0.7 * inch,
            rightMargin=0.7 * inch,
            topMargin=0.7 * inch,
            bottomMargin=0.7 * inch,
        )
        document.build(story)

    def _build_pdf_story_blocks(
        self,
        content: str,
        body_style: ParagraphStyle,
        bullet_style: ParagraphStyle,
        number_style: ParagraphStyle,
    ) -> list[Any]:
        blocks: list[Any] = []
        paragraphs = [block.strip() for block in re.split(r"\n\s*\n", content or "") if block.strip()]
        if not paragraphs:
            blocks.append(Paragraph("", body_style))
            return blocks
        for paragraph in paragraphs:
            lines = [line.rstrip() for line in paragraph.splitlines() if line.strip()]
            if not lines:
                continue
            if all(line.lstrip().startswith(("- ", "* ")) for line in lines):
                for line in lines:
                    blocks.append(Paragraph(f"• {self._markdown_inline_to_rl(line.lstrip()[2:].strip())}", bullet_style))
                continue
            if all(re.match(r"\d+\.\s+", line.lstrip()) for line in lines):
                for line in lines:
                    blocks.append(Paragraph(self._markdown_inline_to_rl(line.lstrip()), number_style))
                continue
            escaped = "<br/>".join(self._markdown_inline_to_rl(line) for line in lines)
            blocks.append(Paragraph(escaped, body_style))
        return blocks

    @staticmethod
    def _markdown_inline_to_rl(text: str) -> str:
        """Convert inline markdown (**bold**, *italic*) to ReportLab XML tags."""
        escaped = xml_escape(text)
        escaped = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", escaped)
        escaped = re.sub(r"\*(.+?)\*", r"<i>\1</i>", escaped)
        escaped = re.sub(r"_(.+?)_", r"<i>\1</i>", escaped)
        return escaped

    async def _translate_with_llm(
        self,
        *,
        session: SessionContext,
        language: str,
        good_examples: list[ExampleSnippet],
    ) -> dict[str, str]:
        payload = await llm_provider.generate_json(
            system_prompt=llm_provider.with_scenario_mcp_context(
                self._build_llm_translation_system_prompt(language),
                session.mcp_servers,
            ),
            user_prompt=self._build_llm_translation_user_prompt(session, language, good_examples),
            temperature=0.1,
            mcp_servers=session.mcp_servers,
        )
        translated_sections: dict[str, str] = {}
        for section in payload.get("sections") or []:
            if not isinstance(section, dict):
                continue
            section_id = section.get("section_id")
            if not isinstance(section_id, str) or not section_id.strip():
                continue
            title = section.get("title")
            content = section.get("content")
            if isinstance(title, str) and title.strip():
                translated_sections[f"{section_id}::title"] = title
            if isinstance(content, str) and content.strip():
                translated_sections[section_id] = content
                translated_sections[f"{section_id}::content"] = content
        return translated_sections

    def _apply_translations_and_structure(
        self,
        *,
        document: DocxDocument,
        session: SessionContext,
        translated_sections: dict[str, str],
        language: str,
    ) -> list[str]:
        warnings: list[str] = []
        original_template = session.original_template_structure or session.template_structure
        original_sections_by_id = {section.id: section for section in original_template.sections}

        for current_section in session.template_structure.sections:
            original_section = original_sections_by_id.get(current_section.id)
            if original_section is None:
                continue
            translated_text = translated_sections.get(f"{current_section.id}::content") or translated_sections.get(current_section.id)
            translated_title = translated_sections.get(f"{current_section.id}::title") or current_section.title
            warning = self._apply_section_translation(document, original_section, translated_text, translated_title)
            if warning:
                warnings.append(f"{language}: {warning}")

        self._rebuild_document_body(document=document, session=session, translated_sections=translated_sections, original_template=original_template)
        return warnings

    def _rebuild_document_body(
        self,
        *,
        document: DocxDocument,
        session: SessionContext,
        translated_sections: dict[str, str],
        original_template: TemplateStructure,
    ) -> None:
        original_sections = list(original_template.sections)
        paragraph_elements = [paragraph._element for paragraph in document.paragraphs]
        if not paragraph_elements:
            return

        ordered_original_sections = [
            section for section in sorted(original_sections, key=self._section_start_index) if self._section_start_index(section) is not None
        ]
        if not ordered_original_sections:
            return

        block_map, first_start, last_end = self._build_original_section_blocks(document, ordered_original_sections)
        body = document._body._element
        prefix = [deepcopy(element) for element in paragraph_elements[:first_start]] if first_start is not None else []
        suffix = [deepcopy(element) for element in paragraph_elements[last_end:]] if last_end is not None else []

        for child in list(body):
            if child.tag.endswith("}p"):
                body.remove(child)

        insert_before = next((child for child in body if child.tag.endswith("}sectPr")), None)
        ordered_elements: list[Any] = []
        ordered_elements.extend(prefix)
        for section in session.template_structure.sections:
            if section.id in block_map:
                ordered_elements.extend(deepcopy(element) for element in block_map[section.id])
                continue
            ordered_elements.extend(self._build_dynamic_section_elements(section, translated_sections))
        ordered_elements.extend(suffix)

        for element in ordered_elements:
            if insert_before is None:
                body.append(element)
            else:
                body.insert(body.index(insert_before), element)

    def _build_original_section_blocks(
        self,
        document: DocxDocument,
        ordered_original_sections: list[TemplateSection],
    ) -> tuple[dict[str, list[Any]], int | None, int | None]:
        block_map: dict[str, list[Any]] = {}
        starts = [self._section_start_index(section) for section in ordered_original_sections]
        valid_starts = [start for start in starts if start is not None]
        if not valid_starts:
            return {}, None, None

        first_start = valid_starts[0]
        last_end = 0
        for index, section in enumerate(ordered_original_sections):
            start = self._section_start_index(section)
            if start is None:
                continue
            next_start = len(document.paragraphs)
            for later in ordered_original_sections[index + 1:]:
                candidate = self._section_start_index(later)
                if candidate is not None:
                    next_start = candidate
                    break
            block_map[section.id] = [document.paragraphs[position]._element for position in range(start, min(next_start, len(document.paragraphs)))]
            last_end = max(last_end, next_start)
        return block_map, first_start, last_end

    @staticmethod
    def _section_start_index(section: TemplateSection) -> int | None:
        if section.heading_paragraph_index is not None:
            return section.heading_paragraph_index
        if section.content_paragraph_indices:
            return min(section.content_paragraph_indices)
        return None

    def _build_dynamic_section_elements(self, section: TemplateSection, translated_sections: dict[str, str]) -> list[Any]:
        heading_text = translated_sections.get(f"{section.id}::title") or section.title
        body_text = translated_sections.get(f"{section.id}::content") or translated_sections.get(section.id) or ""
        scratch_document = DocxDocument()
        scratch_document.add_heading(heading_text, level=max(1, section.level or 1))
        if body_text:
            scratch_document.add_paragraph(body_text)
        return [deepcopy(paragraph._element) for paragraph in scratch_document.paragraphs]

    async def _translate_with_azure(self, *, session: SessionContext, language: str) -> dict[str, str]:
        if not self.settings.azure_translator_key or not self.settings.azure_translator_region:
            raise TranslationProviderConfigurationError(
                "Azure Translator requires AZURE_TRANSLATOR_KEY and AZURE_TRANSLATOR_REGION in backend/.env."
            )

        items = self._collect_translatable_items(session)
        if not items:
            return {}

        headers = {
            "Ocp-Apim-Subscription-Key": self.settings.azure_translator_key,
            "Ocp-Apim-Subscription-Region": self.settings.azure_translator_region,
            "Content-Type": "application/json",
        }
        params: dict[str, str] = {
            "api-version": "3.0",
            "to": self._resolve_language_code(language),
        }
        if self.settings.translation_source_language:
            params["from"] = self.settings.translation_source_language

        try:
            async with httpx.AsyncClient(timeout=self.settings.request_timeout_seconds) as client:
                response = await client.post(
                    f"{self.settings.azure_translator_endpoint.rstrip('/')}/translate",
                    params=params,
                    headers=headers,
                    json=[{"text": item["text"]} for item in items],
                )
            response.raise_for_status()
        except httpx.HTTPStatusError as exc:
            detail = exc.response.text.strip() or str(exc)
            raise TranslationProviderError(f"Azure Translator request failed: {detail}") from exc
        except httpx.HTTPError as exc:
            raise TranslationProviderError("Azure Translator could not be reached with the current configuration.") from exc

        data = response.json()
        if not isinstance(data, list) or len(data) != len(items):
            raise TranslationProviderError("Azure Translator response format was not recognized.")

        return self._build_translation_mapping(items, data, provider_name="Azure Translator")

    async def _translate_with_google(self, *, session: SessionContext, language: str) -> dict[str, str]:
        if not self.settings.google_translate_api_key:
            raise TranslationProviderConfigurationError(
                "Google Translate requires GOOGLE_TRANSLATE_API_KEY in backend/.env."
            )

        items = self._collect_translatable_items(session)
        if not items:
            return {}

        payload: dict[str, Any] = {
            "q": [item["text"] for item in items],
            "target": self._resolve_language_code(language),
            "format": "text",
        }
        if self.settings.translation_source_language:
            payload["source"] = self.settings.translation_source_language

        try:
            async with httpx.AsyncClient(timeout=self.settings.request_timeout_seconds) as client:
                response = await client.post(
                    "https://translation.googleapis.com/language/translate/v2",
                    params={"key": self.settings.google_translate_api_key},
                    json=payload,
                )
            response.raise_for_status()
        except httpx.HTTPStatusError as exc:
            detail = exc.response.text.strip() or str(exc)
            raise TranslationProviderError(f"Google Translate request failed: {detail}") from exc
        except httpx.HTTPError as exc:
            raise TranslationProviderError("Google Translate could not be reached with the current configuration.") from exc

        data = response.json()
        translations = (((data.get("data") or {}).get("translations")) or []) if isinstance(data, dict) else []
        if not isinstance(translations, list) or len(translations) != len(items):
            raise TranslationProviderError("Google Translate response format was not recognized.")

        normalized_payload = [{"translations": [entry]} for entry in translations]
        return self._build_translation_mapping(items, normalized_payload, provider_name="Google Translate")

    def _is_provider_configured(self, provider: TranslationProviderId) -> bool:
        if provider == "llm":
            return bool(self.settings.llm_base_url and self.settings.llm_model)
        if provider == "azure":
            return bool(self.settings.azure_translator_endpoint and self.settings.azure_translator_key and self.settings.azure_translator_region)
        return bool(self.settings.google_translate_api_key)

    def _collect_translatable_items(self, session: SessionContext) -> list[dict[str, str]]:
        items: list[dict[str, str]] = []
        for section in session.draft_state.sections:
            if not section.content.strip():
                continue
            items.append({"section_id": section.section_id, "kind": "title", "text": section.title})
            items.append({"section_id": section.section_id, "kind": "content", "text": section.content})
        return items

    def _build_translation_mapping(self, items: list[dict[str, str]], payload: list[dict[str, Any]], *, provider_name: str) -> dict[str, str]:
        translated_sections: dict[str, str] = {}
        for item, result in zip(items, payload, strict=False):
            translations = result.get("translations") if isinstance(result, dict) else None
            if not isinstance(translations, list) or not translations:
                raise TranslationProviderError(f"{provider_name} did not return a translated text value.")
            translated_text = translations[0].get("text") if isinstance(translations[0], dict) else None
            if not isinstance(translated_text, str):
                raise TranslationProviderError(f"{provider_name} did not return a translated text value.")
            normalized_text = unescape(translated_text)
            section_id = item["section_id"]
            if item["kind"] == "title":
                translated_sections[f"{section_id}::title"] = normalized_text
            else:
                translated_sections[section_id] = normalized_text
                translated_sections[f"{section_id}::content"] = normalized_text
        return translated_sections

    def _build_llm_translation_system_prompt(self, language: str) -> str:
        return (
            "You translate structured business and technical documentation. "
            "Preserve meaning, formatting intent, headings, and professional tone. "
            "Translate both the section titles and the drafted section content into the requested target language. "
            "Return only JSON with a sections array, each containing section_id, title, and content. "
            f"Target language: {language}."
        )

    def _build_llm_translation_user_prompt(
        self,
        session: SessionContext,
        language: str,
        good_examples: list[ExampleSnippet],
    ) -> str:
        examples = "\n\n".join(snippet.content[:700] for snippet in good_examples) or "No good examples uploaded."
        source_sections = "\n\n".join(
            f"Section ID: {section.section_id}\nTitle: {section.title}\nContent:\n{section.content}"
            for section in session.draft_state.sections
            if section.content.strip()
        )
        return (
            f"Translate the drafted document content below into {language}. Do not translate template placeholder text that is not present in the drafted content.\n\n"
            f"Tone guidance from approved examples:\n{examples}\n\n"
            f"Drafted source sections to translate:\n{source_sections}\n\n"
            'Return JSON in this shape: {"sections": [{"section_id": string, "title": string, "content": string}]}'
        )

    def _resolve_language_code(self, language: str) -> str:
        normalized = language.strip()
        if not normalized:
            raise TranslationProviderConfigurationError("A target language is required for translation.")
        language_code = LANGUAGE_CODES.get(normalized.lower())
        if language_code:
            return language_code
        if re.fullmatch(r"[A-Za-z]{2,3}(?:-[A-Za-z]{2,4})?", normalized):
            return normalized
        raise TranslationProviderConfigurationError(
            f"Unsupported target language '{language}'. Use a known language name or ISO language code."
        )

    def _apply_section_translation(
        self,
        document: DocxDocument,
        section: TemplateSection,
        translated_text: str | None,
        translated_title: str | None,
    ) -> str | None:
        if section.heading_paragraph_index is None and not section.content_paragraph_indices:
            heading_text = translated_title or section.title
            if heading_text:
                document.add_heading(heading_text, level=max(1, section.level or 1))
            if translated_text:
                document.add_paragraph(translated_text)
            return None

        if translated_title and section.heading_paragraph_index is not None and section.heading_paragraph_index < len(document.paragraphs):
            document.paragraphs[section.heading_paragraph_index].text = translated_title

        if not translated_text:
            return None

        if not section.content_paragraph_indices:
            return f"Section '{section.title}' has no writable body paragraphs in the template."

        first_index = section.content_paragraph_indices[0]
        if first_index >= len(document.paragraphs):
            return f"Section '{section.title}' points to an invalid paragraph index."

        first_paragraph = document.paragraphs[first_index]
        first_paragraph.text = translated_text

        for redundant_index in section.content_paragraph_indices[1:]:
            if redundant_index < len(document.paragraphs):
                document.paragraphs[redundant_index].text = ""
        return None


translation_service = TranslationService()